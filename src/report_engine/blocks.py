from __future__ import annotations

from pathlib import Path
from typing import Any, Callable, Dict

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from lxml import etree


DEFAULT_STYLE_MAP = {
    "heading_2": "Heading 2",
    "heading_3": "Heading 3",
    "body": "Body Text",
    "caption": "Caption",
    "legend": "Legend",
    "figure_paragraph": "Figure Paragraph",
    "table": "ResearchTable",
    "bullet_list": "List Bullet",
    "numbered_list": "List Number",
    "note": "Note",
    "quote": "Quote",
    "appendix_table": "AppendixTable",
    "checklist": "Checklist",
    "code_block": "CodeBlock",
}


class BlockRenderError(ValueError):
    pass


class BlockRegistry:
    def __init__(self) -> None:
        self._renderers: Dict[str, Callable[..., None]] = {}

    def register(self, block_type: str, renderer_fn: Callable[..., None]) -> None:
        self._renderers[block_type] = renderer_fn

    def render(self, doc: Any, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
        block_type = block["type"]
        renderer = self._renderers.get(block_type)
        if renderer is None:
            raise BlockRenderError(f"Unsupported block type: {block_type}")
        renderer(doc, block, style_map)


def _get_style_name(doc: Any, preferred: str, fallback: str) -> str:
    try:
        style_names = {style.name for style in doc.styles}
    except AttributeError:
        # doc 可能是 _Cell 等不支持 styles 的对象
        return fallback
    return preferred if preferred in style_names else fallback


def _set_table_borders(table: Any) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


MATHML_NS = "http://www.w3.org/1998/Math/MathML"


def _mathml_to_omml(src: etree._Element, parent: etree._Element) -> None:
    """Convert a MathML element tree into OMML (m:oMath children)."""
    tag = etree.QName(src).localname

    if tag in ("math", "mrow"):
        for child in src:
            _mathml_to_omml(child, parent)
    elif tag == "mi":
        r = OxmlElement("m:r")
        rPr = OxmlElement("m:rPr")
        sty = OxmlElement("m:sty")
        sty.set(qn("m:val"), "i")
        rPr.append(sty)
        r.append(rPr)
        t = OxmlElement("m:t")
        t.text = src.text
        r.append(t)
        parent.append(r)
    elif tag in ("mo", "mn"):
        r = OxmlElement("m:r")
        t = OxmlElement("m:t")
        t.text = src.text
        r.append(t)
        parent.append(r)
    elif tag == "msup":
        sSup = OxmlElement("m:sSup")
        e = OxmlElement("m:e")
        _mathml_to_omml(src[0], e)
        sSup.append(e)
        sup = OxmlElement("m:sup")
        _mathml_to_omml(src[1], sup)
        sSup.append(sup)
        parent.append(sSup)
    elif tag == "msub":
        sSub = OxmlElement("m:sSub")
        e = OxmlElement("m:e")
        _mathml_to_omml(src[0], e)
        sSub.append(e)
        sub = OxmlElement("m:sub")
        _mathml_to_omml(src[1], sub)
        sSub.append(sub)
        parent.append(sSub)
    elif tag == "msubsup":
        sSubSup = OxmlElement("m:sSubSup")
        e = OxmlElement("m:e")
        _mathml_to_omml(src[0], e)
        sSubSup.append(e)
        sub = OxmlElement("m:sub")
        _mathml_to_omml(src[1], sub)
        sSubSup.append(sub)
        sup = OxmlElement("m:sup")
        _mathml_to_omml(src[2], sup)
        sSubSup.append(sup)
        parent.append(sSubSup)
    elif tag == "mfrac":
        f = OxmlElement("m:f")
        num = OxmlElement("m:num")
        _mathml_to_omml(src[0], num)
        f.append(num)
        den = OxmlElement("m:den")
        _mathml_to_omml(src[1], den)
        f.append(den)
        parent.append(f)
    else:
        for child in src:
            _mathml_to_omml(child, parent)


def _add_table_block_impl(
    doc: Any,
    block: Dict[str, Any],
    style_map: Dict[str, str],
    default_style: str,
) -> None:
    if block.get("title"):
        caption_style = _get_style_name(doc, style_map["caption"], "Caption")
        doc.add_paragraph(str(block["title"]), style=caption_style)

    headers = block["headers"]
    rows = block["rows"]
    table_style = block.get("style") or default_style
    table_style = _get_style_name(doc, table_style, "Table Grid")

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = table_style

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = "" if header is None else str(header)

    for row in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = "" if value is None else str(value)

    if block.get("force_borders", True):
        _set_table_borders(table)

    body_style = _get_style_name(doc, style_map["body"], "Normal")
    doc.add_paragraph("", style=body_style)


def add_heading_block(doc: Any, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
    level = int(block.get("level", 2))
    if level <= 2:
        style_name = _get_style_name(doc, style_map["heading_2"], "Heading 2")
    else:
        style_name = _get_style_name(doc, style_map["heading_3"], "Heading 3")
    doc.add_paragraph(str(block["text"]), style=style_name)


def add_paragraph_block(doc: Any, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
    style_name = _get_style_name(doc, style_map["body"], "Normal")
    doc.add_paragraph(str(block["text"]), style=style_name)


def add_bullet_list_block(doc: Any, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
    style_name = _get_style_name(doc, style_map["bullet_list"], style_map["body"])
    for item in block["items"]:
        p = doc.add_paragraph(style=style_name)
        p.add_run(str(item))


def add_numbered_list_block(doc: Any, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
    style_name = _get_style_name(doc, style_map["numbered_list"], style_map["body"])
    for item in block["items"]:
        p = doc.add_paragraph(style=style_name)
        p.add_run(str(item))


def add_page_break_block(doc: Any, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_table_block(doc: Any, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
    _add_table_block_impl(doc, block, style_map, style_map["table"])


def add_image_block(doc: Any, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
    image_path = Path(block["path"])

    figure_style = _get_style_name(doc, style_map["figure_paragraph"], style_map["body"])
    p = doc.add_paragraph(style=figure_style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if image_path.exists():
        run = p.add_run()
        width_cm = block.get("width_cm")
        if width_cm is not None:
            run.add_picture(str(image_path), width=Cm(float(width_cm)))
        else:
            run.add_picture(str(image_path))
    else:
        p.add_run(f"[图片缺失：{image_path}]")

    if block.get("caption"):
        caption_style = _get_style_name(doc, style_map["caption"], "Caption")
        cp = doc.add_paragraph(str(block["caption"]), style=caption_style)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if block.get("legend"):
        legend_style = _get_style_name(doc, style_map["legend"], style_map["body"])
        lp = doc.add_paragraph(str(block["legend"]), style=legend_style)
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    body_style = _get_style_name(doc, style_map["body"], "Normal")
    doc.add_paragraph("", style=body_style)


def add_rich_paragraph_block(doc: Any, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
    style_name = _get_style_name(doc, style_map["body"], "Normal")
    p = doc.add_paragraph(style=style_name)
    for seg in block["segments"]:
        run = p.add_run(str(seg.get("text", "")))
        if seg.get("bold"):
            run.bold = True
        if seg.get("italic"):
            run.italic = True
        if seg.get("sub"):
            rpr = run._element.get_or_add_rPr()
            vert_align = OxmlElement("w:vertAlign")
            vert_align.set(qn("w:val"), "subscript")
            rpr.append(vert_align)
        if seg.get("sup"):
            rpr = run._element.get_or_add_rPr()
            vert_align = OxmlElement("w:vertAlign")
            vert_align.set(qn("w:val"), "superscript")
            rpr.append(vert_align)


def add_note_block(doc: Any, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
    style_name = _get_style_name(doc, style_map.get("note", "Note"), style_map["body"])
    p = doc.add_paragraph(style=style_name)
    prefix_run = p.add_run("注：")
    prefix_run.bold = True
    p.add_run(str(block["text"]))


def add_quote_block(doc: Any, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
    quote_style = _get_style_name(doc, style_map.get("quote", "Quote"), style_map["body"])
    doc.add_paragraph(str(block["text"]), style=quote_style)
    if block.get("source"):
        source_style = _get_style_name(doc, style_map["body"], "Normal")
        sp = doc.add_paragraph(str(block["source"]), style=source_style)
        sp.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_two_images_row_block(doc: Any, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
    images = block["images"]
    if len(images) != 2:
        raise BlockRenderError(f"two_images_row requires exactly 2 images, got {len(images)}")

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 移除边框
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "none")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "auto")
        borders.append(element)
    tbl_pr.append(borders)

    figure_style = _get_style_name(doc, style_map["figure_paragraph"], style_map["body"])
    caption_style = _get_style_name(doc, style_map["caption"], "Caption")

    for i, img in enumerate(images):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        figure_style_name = _get_style_name(doc, figure_style, style_map["body"])
        p.style = doc.styles[figure_style_name]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        image_path = Path(img["path"])
        if image_path.exists():
            run = p.add_run()
            width_cm = img.get("width_cm")
            if width_cm is not None:
                run.add_picture(str(image_path), width=Cm(float(width_cm)))
            else:
                run.add_picture(str(image_path))
        else:
            p.add_run(f"[图片缺失：{image_path}]")

        if img.get("caption"):
            cp = cell.add_paragraph(str(img["caption"]))
            caption_style_name = _get_style_name(doc, caption_style, style_map["body"])
            cp.style = doc.styles[caption_style_name]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    body_style = _get_style_name(doc, style_map["body"], "Normal")
    doc.add_paragraph("", style=body_style)


def add_appendix_table_block(doc: Any, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
    default_style = style_map.get("appendix_table", style_map["table"])
    _add_table_block_impl(doc, block, style_map, default_style)


def add_checklist_block(doc: Any, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
    style_name = _get_style_name(doc, style_map.get("checklist", "Checklist"), style_map.get("bullet_list", style_map["body"]))
    for item in block["items"]:
        p = doc.add_paragraph(style=style_name)
        prefix = "☑" if item.get("checked", False) else "☐"
        p.add_run(f"{prefix} {str(item['text'])}")


def add_horizontal_rule_block(doc: Any, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_toc_placeholder_block(doc: Any, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
    title = block.get("title", "目录")
    title_style = _get_style_name(doc, style_map.get("heading_2", "Heading 2"), "Heading 2")
    doc.add_paragraph(str(title), style=title_style)

    # 插入 TOC 域代码
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fldChar_begin)

    run2 = p.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._element.append(instrText)

    run3 = p.add_run()
    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")
    run3._element.append(fldChar_separate)

    run4 = p.add_run("[请右键更新域以生成目录]")

    run5 = p.add_run()
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run5._element.append(fldChar_end)


def add_code_block_block(doc: Any, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
    code = str(block["code"])
    lines = code.split("\n")
    style_name = _get_style_name(doc, style_map.get("code_block", "CodeBlock"), style_map["body"])

    for line in lines:
        p = doc.add_paragraph(style=style_name)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Cm(0.28)  # ~8pt
        # 灰色底纹
        rPr = run._element.get_or_add_rPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        rPr.append(shd)

    # 语言标注（可选）
    if block.get("language"):
        lang_style = _get_style_name(doc, style_map["body"], "Normal")
        lp = doc.add_paragraph(style=lang_style)
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lr = lp.add_run(str(block["language"]))
        lr.font.size = Cm(0.21)  # ~6pt
        rPr = lr._element.get_or_add_rPr()
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "808080")
        rPr.append(color)


def add_formula_block(doc: Any, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
    latex = str(block["latex"])

    # 方案 1: LaTeX → MathML → OMML
    omml_inserted = False
    try:
        from latex2mathml.converter import convert as latex_to_mathml

        mathml_str = latex_to_mathml(latex)
        mathml_root = etree.fromstring(mathml_str.encode("utf-8"))

        style_name = _get_style_name(doc, style_map["body"], "Normal")
        p = doc.add_paragraph(style=style_name)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        oMath = OxmlElement("m:oMath")
        _mathml_to_omml(mathml_root, oMath)
        p._element.append(oMath)
        omml_inserted = True
    except Exception:
        pass

    # 方案 2: LaTeX → 图片（降级）
    if not omml_inserted:
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt
            from io import BytesIO

            fig, ax = plt.subplots(figsize=(0.01, 0.01))
            ax.axis("off")
            text = ax.text(0.5, 0.5, f"${latex}$", fontsize=14, ha="center", va="center")
            buf = BytesIO()
            fig.savefig(buf, format="png", bbox_inches="tight", dpi=150, transparent=True)
            plt.close(fig)
            buf.seek(0)

            style_name = _get_style_name(doc, style_map["body"], "Normal")
            p = doc.add_paragraph(style=style_name)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(buf, width=Cm(8))
            omml_inserted = True
        except Exception:
            pass

    # 方案 3: 纯文本降级
    if not omml_inserted:
        style_name = _get_style_name(doc, style_map.get("code_block", "CodeBlock"), style_map["body"])
        p = doc.add_paragraph(style=style_name)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(latex)
        run.font.name = "Courier New"

    # caption（可选）
    if block.get("caption"):
        caption_style = _get_style_name(doc, style_map["caption"], "Caption")
        cp = doc.add_paragraph(str(block["caption"]), style=caption_style)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _has_cjk(text: str) -> bool:
    """检测文本是否包含 CJK 字符。"""
    for ch in text:
        cp = ord(ch)
        if (
            0x4E00 <= cp <= 0x9FFF  # CJK Unified Ideographs
            or 0x3400 <= cp <= 0x4DBF  # CJK Extension A
            or 0x3000 <= cp <= 0x303F  # CJK Symbols and Punctuation
            or 0xFF00 <= cp <= 0xFFEF  # Halfwidth and Fullwidth Forms
            or 0x3040 <= cp <= 0x309F  # Hiragana
            or 0x30A0 <= cp <= 0x30FF  # Katakana
            or 0xAC00 <= cp <= 0xD7AF  # Hangul Syllables
        ):
            return True
    return False


def add_ascii_diagram_block(doc: Any, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
    """将 ASCII 文本渲染为图片插入文档。

    策略：
    - 纯 ASCII（无 CJK）：用 Pillow + 等宽字体渲染为图片
    - 含 CJK 字符：用 Pillow 双字体渲染（ASCII 等宽字体 + CJK 字体），
      CJK 字符占 2 列宽度，通过网格对齐保持 ASCII 图结构
    - Pillow 不可用或字体缺失时降级为 Word 内等宽文本表格渲染
    """
    ascii_text = str(block["ascii"])
    lines = ascii_text.split("\n")

    # 样式参数
    font_size = int(block.get("font_size", 14))
    padding = int(block.get("padding", 20))
    bg_color = block.get("bg_color", "#F8F8F8")
    fg_color = block.get("fg_color", "#333333")

    # 检测是否含 CJK 字符
    contains_cjk = _has_cjk(ascii_text)

    if not contains_cjk:
        # 纯 ASCII：用 Pillow 渲染为图片
        try:
            from PIL import Image, ImageDraw, ImageFont
            from io import BytesIO

            # 尝试加载等宽字体
            font_paths = [
                "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf",
                "/usr/share/fonts/truetype/liberation/LiberationMono-Regular.ttf",
                "/usr/share/fonts/truetype/courier/Courier_New.ttf",
                "/System/Library/Fonts/Courier.ttc",
                "C:/Windows/fonts/cour.ttf",
            ]
            font = None
            for fp in font_paths:
                try:
                    font = ImageFont.truetype(fp, font_size)
                    break
                except Exception:
                    continue
            if font is None:
                font = ImageFont.load_default()

            # 计算尺寸（逐字符测量，处理不等宽情况）
            max_width = 0
            total_height = 0
            line_heights = []
            line_widths = []

            for line in lines:
                line_width = 0
                line_height = 0
                for ch in line:
                    bbox = font.getbbox(ch)
                    if bbox:
                        ch_w = bbox[2] - bbox[0]
                        ch_h = bbox[3] - bbox[1]
                        line_width += ch_w
                        line_height = max(line_height, ch_h)
                line_widths.append(line_width)
                line_heights.append(line_height)
                max_width = max(max_width, line_width)
                total_height += int(line_height * 1.3) if line_height else font_size

            line_spacing = int((line_heights[0] if line_heights else font_size) * 1.3)
            if line_spacing == 0:
                line_spacing = font_size + 4

            img_width = max_width + padding * 2
            img_height = total_height + padding * 2

            # 创建图片
            img = Image.new("RGB", (img_width, img_height), bg_color)
            draw = ImageDraw.Draw(img)

            y = padding
            for line in lines:
                x = padding
                for ch in line:
                    draw.text((x, y), ch, font=font, fill=fg_color)
                    bbox = font.getbbox(ch)
                    if bbox:
                        x += bbox[2] - bbox[0]
                y += line_spacing

            # 保存到内存缓冲区
            buf = BytesIO()
            img.save(buf, format="PNG")
            buf.seek(0)

            # 插入到 Word
            figure_style = _get_style_name(
                doc, style_map.get("figure_paragraph", style_map["body"]), style_map["body"]
            )
            p = doc.add_paragraph(style=figure_style)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            width_cm = block.get("width_cm")
            if width_cm is not None:
                run.add_picture(buf, width=Cm(float(width_cm)))
            else:
                run.add_picture(buf)

            # caption
            if block.get("caption"):
                caption_style = _get_style_name(doc, style_map["caption"], "Caption")
                cp = doc.add_paragraph(str(block["caption"]), style=caption_style)
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

            body_style = _get_style_name(doc, style_map["body"], "Normal")
            doc.add_paragraph("", style=body_style)
            return

        except Exception:
            # Pillow 失败则降级为文本渲染
            pass

    # 含 CJK：用 Pillow 双字体渲染（ASCII 字体 + CJK 字体），失败则降级为 Word 文本
    try:
        from PIL import Image, ImageDraw, ImageFont
        from io import BytesIO

        # 加载 ASCII 等宽字体
        ascii_font_paths = [
            "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationMono-Regular.ttf",
            "/usr/share/fonts/truetype/courier/Courier_New.ttf",
            "/System/Library/Fonts/Courier.ttc",
            "C:/Windows/fonts/cour.ttf",
        ]
        ascii_font = None
        for fp in ascii_font_paths:
            try:
                ascii_font = ImageFont.truetype(fp, font_size)
                break
            except Exception:
                continue
        if ascii_font is None:
            ascii_font = ImageFont.load_default()

        # 加载 CJK 字体
        cjk_font_paths = [
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/System/Library/Fonts/PingFang.ttc",
            "C:/Windows/fonts/msyh.ttc",
            "C:/Windows/fonts/simsun.ttc",
        ]
        cjk_font = None
        for fp in cjk_font_paths:
            try:
                cjk_font = ImageFont.truetype(fp, font_size)
                break
            except Exception:
                continue

        if cjk_font is None:
            raise RuntimeError("No CJK font found")

        def _measure(font: ImageFont.FreeTypeFont, ch: str) -> int:
            bbox = font.getbbox(ch)
            return bbox[2] - bbox[0] if bbox else font_size

        # 计算列宽：确保 CJK（占 2 列）能放下
        max_ascii_w = 1
        max_cjk_w = 1
        for line in lines:
            for ch in line:
                if _has_cjk(ch):
                    max_cjk_w = max(max_cjk_w, _measure(cjk_font, ch))
                else:
                    max_ascii_w = max(max_ascii_w, _measure(ascii_font, ch))
        col_width = max(max_ascii_w, (max_cjk_w + 1) // 2 + 1)

        # 计算图片尺寸
        max_cols = 0
        for line in lines:
            cols = sum(2 if _has_cjk(ch) else 1 for ch in line)
            max_cols = max(max_cols, cols)

        line_height = int(font_size * 1.5)
        img_width = max_cols * col_width + padding * 2
        img_height = len(lines) * line_height + padding * 2

        img = Image.new("RGB", (img_width, img_height), bg_color)
        draw = ImageDraw.Draw(img)

        for row, line in enumerate(lines):
            y = padding + row * line_height
            col = 0
            for ch in line:
                x = padding + col * col_width
                if _has_cjk(ch):
                    w = _measure(cjk_font, ch)
                    cx = x + col_width - w // 2
                    draw.text((cx, y), ch, font=cjk_font, fill=fg_color)
                    col += 2
                else:
                    w = _measure(ascii_font, ch)
                    cx = x + (col_width - w) // 2
                    draw.text((cx, y), ch, font=ascii_font, fill=fg_color)
                    col += 1

        buf = BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)

        figure_style = _get_style_name(
            doc, style_map.get("figure_paragraph", style_map["body"]), style_map["body"]
        )
        p = doc.add_paragraph(style=figure_style)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        width_cm = block.get("width_cm")
        if width_cm is not None:
            run.add_picture(buf, width=Cm(float(width_cm)))
        else:
            run.add_picture(buf)

        if block.get("caption"):
            caption_style = _get_style_name(doc, style_map["caption"], "Caption")
            cp = doc.add_paragraph(str(block["caption"]), style=caption_style)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        body_style = _get_style_name(doc, style_map["body"], "Normal")
        doc.add_paragraph("", style=body_style)
        return

    except Exception:
        # Pillow 双字体渲染失败，降级为 Word 文本渲染
        _render_ascii_as_text(doc, block, style_map, lines)


def _render_ascii_as_text(
    doc: Any, block: Dict[str, Any], style_map: Dict[str, str], lines: List[str]
) -> None:
    """在 Word 中用等宽文本段落渲染 ASCII 图。"""
    from docx.oxml import OxmlElement
    from docx.shared import Pt

    font_size = int(block.get("font_size", 10))
    bg = block.get("bg_color", "#F8F8F8").lstrip("#")

    # 创建一个居中的容器段落
    figure_style = _get_style_name(
        doc, style_map.get("figure_paragraph", style_map["body"]), style_map["body"]
    )
    container = doc.add_paragraph(style=figure_style)
    container.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加灰色底纹的表格来包裹 ASCII 文本（保持对齐且美观）
    table = doc.add_table(rows=len(lines), cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 移除表格边框，设置背景色
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

    # 设置表格宽度（使用 width_cm 或自动）
    width_cm = block.get("width_cm")
    if width_cm is not None:
        tbl_w = OxmlElement("w:tblW")
        tbl_w.set(qn("w:w"), str(int(float(width_cm) * 567)))
        tbl_w.set(qn("w:type"), "dxa")
        tbl_pr.append(tbl_w)

    # 设置表格背景色
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), bg)
    tbl_pr.append(shading)

    # 移除边框
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "none")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "auto")
        borders.append(element)
    tbl_pr.append(borders)

    for idx, line in enumerate(lines):
        cell = table.cell(idx, 0)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        run.font.size = Pt(font_size)

        # 手动完整设置 rFonts，避免样式冲突
        rpr = run._element.get_or_add_rPr()
        for existing in list(rpr.findall(qn("w:rFonts"))):
            rpr.remove(existing)
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Courier New")
        rFonts.set(qn("w:hAnsi"), "Courier New")
        rFonts.set(qn("w:eastAsia"), "SimSun")
        rFonts.set(qn("w:cs"), "Courier New")
        rpr.insert(0, rFonts)

        # 设置语言属性，确保 Word 识别为中文并启用东亚字体
        for existing in list(rpr.findall(qn("w:lang"))):
            rpr.remove(existing)
        lang = OxmlElement("w:lang")
        lang.set(qn("w:val"), "en-US")
        lang.set(qn("w:eastAsia"), "zh-CN")
        rpr.append(lang)

    # caption（可选）
    if block.get("caption"):
        caption_style = _get_style_name(doc, style_map["caption"], "Caption")
        cp = doc.add_paragraph(str(block["caption"]), style=caption_style)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    body_style = _get_style_name(doc, style_map["body"], "Normal")
    doc.add_paragraph("", style=body_style)


def add_columns_block(doc: Any, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
    count = int(block["count"])
    columns = block["columns"]
    if len(columns) != count:
        raise BlockRenderError(f"columns: expected {count} columns, got {len(columns)}")

    table = doc.add_table(rows=1, cols=count)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 移除边框
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "none")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "auto")
        borders.append(element)
    tbl_pr.append(borders)

    # 设置列间距（gap_cm）
    gap_cm = block.get("gap_cm")
    if gap_cm is not None:
        tbl_cell_spacing = OxmlElement("w:tblCellSpacing")
        tbl_cell_spacing.set(qn("w:type"), "dxa")
        gap_dxa = int(float(gap_cm) * 567)
        tbl_cell_spacing.set(qn("w:w"), str(gap_dxa))
        tbl_pr.append(tbl_cell_spacing)

    registry = create_default_registry()

    for i, col_blocks in enumerate(columns):
        cell = table.cell(0, i)
        cell.text = ""
        for b in col_blocks:
            registry.render(cell, b, style_map)

    body_style = _get_style_name(doc, style_map["body"], "Normal")
    doc.add_paragraph("", style=body_style)


def create_default_registry() -> BlockRegistry:
    registry = BlockRegistry()
    registry.register("heading", add_heading_block)
    registry.register("paragraph", add_paragraph_block)
    registry.register("bullet_list", add_bullet_list_block)
    registry.register("numbered_list", add_numbered_list_block)
    registry.register("table", add_table_block)
    registry.register("image", add_image_block)
    registry.register("page_break", add_page_break_block)
    registry.register("rich_paragraph", add_rich_paragraph_block)
    registry.register("note", add_note_block)
    registry.register("quote", add_quote_block)
    registry.register("two_images_row", add_two_images_row_block)
    registry.register("appendix_table", add_appendix_table_block)
    registry.register("checklist", add_checklist_block)
    registry.register("horizontal_rule", add_horizontal_rule_block)
    registry.register("toc_placeholder", add_toc_placeholder_block)
    registry.register("code_block", add_code_block_block)
    registry.register("formula", add_formula_block)
    registry.register("columns", add_columns_block)
    registry.register("ascii_diagram", add_ascii_diagram_block)
    return registry
