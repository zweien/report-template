#!/usr/bin/env python3
"""分析 report-engine 模板，输出结构信息供 report-generator skill 使用。

用法:
    python analyze_template.py --template path/to/template.docx
"""

import argparse
import json
from pathlib import Path
from docx import Document


def analyze_template(template_path: str) -> dict:
    """分析模板结构。"""
    doc = Document(template_path)

    # 提取所有文本
    all_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            all_text.append(para.text.strip())

    full_text = "\n".join(all_text)

    # 识别占位符
    import re
    scalars = re.findall(r"\{\{(\w+)\}\}", full_text)
    subdocs = re.findall(r"\{\{p\s+(\w+)\s*\}\}", full_text)
    flags = re.findall(r"\{%p\s+if\s+(\w+)\s*%\}", full_text)

    # 识别样式
    styles = set()
    for para in doc.paragraphs:
        if para.style:
            styles.add(para.style.name)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.style:
                        styles.add(para.style.name)

    # 识别标题结构
    headings = []
    for para in doc.paragraphs:
        if para.style and para.style.name.startswith("Heading"):
            headings.append({
                "level": int(para.style.name.replace("Heading ", "")) if para.style.name[-1].isdigit() else 2,
                "text": para.text.strip(),
            })

    return {
        "path": template_path,
        "scalar_placeholders": sorted(set(scalars)),
        "subdoc_placeholders": sorted(set(subdocs)),
        "conditional_flags": sorted(set(flags)),
        "styles": sorted(styles),
        "headings": headings,
        "paragraphs_count": len(doc.paragraphs),
        "tables_count": len(doc.tables),
    }


def main():
    parser = argparse.ArgumentParser(description="分析 report-engine 模板")
    parser.add_argument("--template", required=True, help="模板文件路径")
    parser.add_argument("--json", action="store_true", help="输出 JSON 格式")
    args = parser.parse_args()

    result = analyze_template(args.template)

    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print(f"模板: {result['path']}")
        print(f"段落数: {result['paragraphs_count']}")
        print(f"表格数: {result['tables_count']}")
        print(f"\n标量占位符: {', '.join(result['scalar_placeholders']) or '无'}")
        print(f"Subdoc 占位符: {', '.join(result['subdoc_placeholders']) or '无'}")
        print(f"条件开关: {', '.join(result['conditional_flags']) or '无'}")
        print(f"\n样式: {', '.join(result['styles'])}")
        print(f"\n标题结构:")
        for h in result["headings"]:
            indent = "  " * (h["level"] - 1)
            print(f"  {indent}Heading {h['level']}: {h['text']}")


if __name__ == "__main__":
    main()
