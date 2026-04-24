#!/usr/bin/env python3
"""
解析模板中的 [[FORMAT: ...]] 格式要求注释，应用到模板样式后去除这些注释段落。

用法:
    python apply_template_formats.py \
        --input templates/draft.docx \
        --output templates/clean.docx \
        [--dry-run]

格式要求语法:
    [[FORMAT: 要求名=值, 要求名=值, ...]]

支持的要求:
    - 字体: 宋体、黑体、楷体等
    - 字号: 初号~八号、小四、四号等
    - 颜色: #RRGGBB 或颜色名
    - 行距: 1.0倍~3.0倍
    - 段前/段后: 0pt~72pt
    - 页边距上/下/左/右: cm 单位
    - 对齐: 左对齐、居中、右对齐、两端对齐
    - 粗体/斜体: true/false
"""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional

import docx.oxml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


FORMAT_PREFIX = "[[FORMAT:"
FORMAT_SUFFIX = "]]"

# 字号映射：名称 -> 磅值
FONT_SIZE_MAP = {
    "初号": 42,
    "小初": 36,
    "一号": 26,
    "小一": 24,
    "二号": 22,
    "小二": 18,
    "三号": 16,
    "小三": 15,
    "四号": 14,
    "小四": 12,
    "五号": 10.5,
    "小五": 9,
    "六号": 7.5,
    "小六": 6.5,
    "七号": 5.5,
    "八号": 5,
}

# 颜色名映射
COLOR_MAP = {
    "黑色": "000000",
    "红色": "FF0000",
    "蓝色": "0000FF",
    "绿色": "008000",
    "黄色": "FFFF00",
    "橙色": "FFA500",
    "紫色": "800080",
    "灰色": "808080",
    "白色": "FFFFFF",
}

# 对齐方式映射
ALIGN_MAP = {
    "左对齐": WD_ALIGN_PARAGRAPH.LEFT,
    "居中": WD_ALIGN_PARAGRAPH.CENTER,
    "右对齐": WD_ALIGN_PARAGRAPH.RIGHT,
    "两端对齐": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


@dataclass
class FormatRequirement:
    """单个格式要求"""
    target: str  # 目标样式名或"页面"
    properties: Dict[str, str]


def _parse_font_size(value: str) -> Optional[Pt]:
    """解析字号为 Pt"""
    value = value.strip()
    if value in FONT_SIZE_MAP:
        return Pt(FONT_SIZE_MAP[value])
    # 尝试解析数字（磅值）
    match = re.match(r"(\d+(?:\.\d+)?)", value)
    if match:
        return Pt(float(match.group(1)))
    return None


def _parse_color(value: str) -> Optional[RGBColor]:
    """解析颜色为 RGBColor"""
    value = value.strip()
    if value in COLOR_MAP:
        return RGBColor.from_string(COLOR_MAP[value])
    # 尝试解析 #RRGGBB
    match = re.match(r"#?([0-9A-Fa-f]{6})", value)
    if match:
        return RGBColor.from_string(match.group(1))
    return None


def _parse_bool(value: str) -> Optional[bool]:
    """解析布尔值"""
    value = value.strip().lower()
    if value in ("true", "是", "yes", "1"):
        return True
    if value in ("false", "否", "no", "0"):
        return False
    return None


def _parse_line_spacing(value: str) -> Optional[float]:
    """解析行距倍数"""
    match = re.match(r"(\d+(?:\.\d+)?)", value.strip())
    if match:
        return float(match.group(1))
    return None


def _parse_pt(value: str) -> Optional[Pt]:
    """解析磅值"""
    match = re.match(r"(\d+(?:\.\d+)?)", value.strip())
    if match:
        return Pt(float(match.group(1)))
    return None


def _parse_cm(value: str) -> Optional[Cm]:
    """解析厘米值"""
    match = re.match(r"(\d+(?:\.\d+)?)", value.strip())
    if match:
        return Cm(float(match.group(1)))
    return None


def _extract_number_unit(text: str) -> tuple[str, str]:
    """从文本中提取数字和单位，如 '0.76厘米' -> ('0.76', '厘米')"""
    match = re.search(r"(\d+(?:\.\d+)?)\s*([a-zA-Z\u4e00-\u9fff]+)", text)
    if match:
        return match.group(1), match.group(2)
    return "", ""


def parse_natural_language_format(text: str) -> tuple[Optional[str], Dict[str, str]]:
    """从自然语言描述中提取格式要求。

    支持的描述模式:
    - 字体均为五号宋体 / 字体为黑体 / 使用楷体
    - 字号为三号 / 使用小四字号
    - 版式采用两端对齐 / 左对齐 / 居中
    - 首行缩进0.76厘米 / 首行缩进2字符
    - 段前0磅 / 段后12磅
    - 行距1.5倍 / 行距最小值18磅 / 行距固定值20磅
    - 粗体 / 不加粗
    - 颜色为红色 / #FF0000
    """
    properties: Dict[str, str] = {}
    target = None

    # 推断目标样式
    if "正文" in text or "内容" in text:
        target = "Body Text"
    elif "标题" in text:
        target = "Heading 2"
    elif "表格" in text:
        target = "ResearchTable"
    elif "代码" in text:
        target = "CodeBlock"
    elif "页面" in text or "版式" in text:
        target = "页面"

    # 字体 + 字号组合: "五号宋体" / "小四宋体" / "三号黑体"
    font_size_pattern = re.compile(
        r"([初小一二三四五六七八号]+)\s*([宋黑楷仿隶微软雅圆宋体]+)"
    )
    m = font_size_pattern.search(text)
    if m:
        size_name, font_name = m.group(1), m.group(2)
        properties["字号"] = size_name
        properties["字体"] = font_name
    else:
        # 单独提取字体: "字体为宋体" / "使用黑体" / "楷体"
        m = re.search(r"(?:字体|使用|采用)\s*(?:为|的|是)?\s*([宋黑楷仿隶圆雅微软]+(?:体|雅黑)?)", text)
        if m:
            properties["字体"] = m.group(1)
        # 单独提取字号: "字号为五号" / "使用三号字号"
        m = re.search(r"(?:字号|大小)\s*(?:为|的|是)?\s*([初小一二三四五六七八号]+)", text)
        if m:
            properties["字号"] = m.group(1)

    # 对齐方式
    align_keywords = {
        "两端对齐": "两端对齐",
        "左对齐": "左对齐",
        "右对齐": "右对齐",
        "居中": "居中",
        "居中对齐": "居中",
    }
    for kw, val in align_keywords.items():
        if kw in text:
            properties["对齐"] = val
            break

    # 支持中文/英文引号的辅助模式
    QUOTE_PAT = r"[\"\"'\"'\"'‘’]?"

    # 首行缩进
    m = re.search(
        rf"首行缩进\s*{QUOTE_PAT}(\d+(?:\.\d+)?)\s*([厘米cm字符字]+){QUOTE_PAT}", text
    )
    if m:
        num, unit = m.group(1), m.group(2)
        if "厘米" in unit or "cm" in unit:
            properties["首行缩进"] = f"{num}cm"
        elif "字符" in unit or "字" in unit:
            # Word 中首行缩进 2 字符 ≈ 0.74cm (以五号字 10.5pt 计算)
            properties["首行缩进"] = f"{float(num) * 0.37}cm"

    # 段前
    m = re.search(rf"段前\s*{QUOTE_PAT}(\d+(?:\.\d+)?)\s*([磅pt磅]+){QUOTE_PAT}", text)
    if m:
        properties["段前"] = f"{m.group(1)}pt"

    # 段后
    m = re.search(rf"段后\s*{QUOTE_PAT}(\d+(?:\.\d+)?)\s*([磅pt磅]+){QUOTE_PAT}", text)
    if m:
        properties["段后"] = f"{m.group(1)}pt"

    # 行距 - 三种模式
    # 倍数: "行距1.5倍" / "行距为 2 倍"
    m = re.search(r"行距\s*(?:为|采用)?\s*(\d+(?:\.\d+)?)\s*倍", text)
    if m:
        properties["行距"] = m.group(1)
    else:
        # 最小值: "行距最小值18磅" 或 "行距"最小值18磅""
        m = re.search(
            rf"行距\s*{QUOTE_PAT}最小值\s*(\d+(?:\.\d+)?)\s*([磅pt]+){QUOTE_PAT}", text
        )
        if m:
            properties["行距最小值"] = f"{m.group(1)}pt"
        else:
            # 固定值: "行距固定值20磅"
            m = re.search(
                rf"行距\s*{QUOTE_PAT}固定值\s*(\d+(?:\.\d+)?)\s*([磅pt]+){QUOTE_PAT}", text
            )
            if m:
                properties["行距固定值"] = f"{m.group(1)}pt"

    # 粗体
    if re.search(r"(?:不加粗|取消粗体|非粗体)", text):
        properties["粗体"] = "false"
    elif re.search(r"(?:粗体|加粗)", text):
        properties["粗体"] = "true"

    # 斜体
    if re.search(r"(?:不斜体|取消斜体)", text):
        properties["斜体"] = "false"
    elif re.search(r"(?:斜体|倾斜)", text):
        properties["斜体"] = "true"

    # 颜色
    m = re.search(
        r"(?:颜色|色)\s*(?:为|的|是)?\s*(#[0-9A-Fa-f]{6}|[红蓝绿黄橙紫黑白灰色]+)", text
    )
    if m:
        properties["颜色"] = m.group(1)

    # 页面设置
    for direction, key in [("上", "页边距上"), ("下", "页边距下"), ("左", "页边距左"), ("右", "页边距右")]:
        m = re.search(
            rf"页边距[{direction}]\s*{QUOTE_PAT}(\d+(?:\.\d+)?)\s*([厘米cm]+){QUOTE_PAT}", text
        )
        if m:
            properties[key] = f"{m.group(1)}cm"

    return target, properties


def parse_format_comment(text: str) -> Optional[FormatRequirement]:
    """解析 [[FORMAT: ...]] 注释文本"""
    text = text.strip()
    if not (text.startswith(FORMAT_PREFIX) and text.endswith(FORMAT_SUFFIX)):
        return None

    content = text[len(FORMAT_PREFIX) : -len(FORMAT_SUFFIX)].strip()
    if not content:
        return None

    # 检测是否是键值对格式（包含 =）还是自然语言描述
    if "=" in content and not re.search(r"[^=]\s+为\s+[^=]", content):
        # 键值对格式: "目标样式: 属性=值, 属性=值"
        target = None
        if ":" in content:
            target_part, props_part = content.split(":", 1)
            target = target_part.strip()
            content = props_part.strip()
        else:
            # 从属性名智能推断目标样式
            for key in content.split(","):
                key_name = key.split("=")[0].strip().lower() if "=" in key else ""
                if "标题" in key_name:
                    target = "Heading 2"
                    break
                elif "正文" in key_name:
                    target = "Body Text"
                    break
                elif "表格" in key_name or "表头" in key_name:
                    target = "ResearchTable"
                    break
                elif "图片" in key_name or "图题" in key_name:
                    target = "Figure Paragraph"
                    break
                elif "代码" in key_name:
                    target = "CodeBlock"
                    break
                elif "引用" in key_name:
                    target = "Quote"
                    break
                elif "注释" in key_name or "注" in key_name:
                    target = "Note"
                    break
                elif "页面" in key_name or "页边距" in key_name or "纸张" in key_name:
                    target = "页面"
                    break
            if target is None:
                target = "Body Text"

        properties = {}
        for part in re.split(r",(?![^\(]*\))", content):
            part = part.strip()
            if "=" in part:
                k, v = part.split("=", 1)
                properties[k.strip()] = v.strip()

        return FormatRequirement(target=target, properties=properties)

    # 自然语言描述格式
    target, properties = parse_natural_language_format(content)
    if target is None:
        target = "Body Text"

    return FormatRequirement(target=target, properties=properties)


def apply_paragraph_style_format(
    doc: Document, style_name: str, properties: Dict[str, str]
) -> List[str]:
    """应用段落样式格式，返回应用的变更描述"""
    changes = []

    # 查找或创建样式
    try:
        style = doc.styles[style_name]
    except KeyError:
        from docx.enum.style import WD_STYLE_TYPE
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        changes.append(f"创建样式: {style_name}")

    pf = style.paragraph_format
    font = style.font

    for key, value in properties.items():
        key_lower = key.lower()

        if "字体" in key_lower or key_lower == "font":
            font.name = value
            # 同时设置东亚字体
            rpr = style.element.get_or_add_rPr()
            rFonts = rpr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = docx.oxml.OxmlElement("w:rFonts")
                rpr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), value)
            changes.append(f"{style_name}.字体 = {value}")

        elif key_lower in ("字号", "size", "fontsize"):
            pt = _parse_font_size(value)
            if pt:
                font.size = pt
                changes.append(f"{style_name}.字号 = {value} ({pt.pt}pt)")

        elif key_lower in ("颜色", "color", "fontcolor"):
            color = _parse_color(value)
            if color:
                font.color.rgb = color
                changes.append(f"{style_name}.颜色 = {value}")

        elif key_lower in ("粗体", "bold"):
            b = _parse_bool(value)
            if b is not None:
                font.bold = b
                changes.append(f"{style_name}.粗体 = {b}")

        elif key_lower in ("斜体", "italic"):
            b = _parse_bool(value)
            if b is not None:
                font.italic = b
                changes.append(f"{style_name}.斜体 = {b}")

        elif key_lower in ("行距", "linespacing"):
            ls = _parse_line_spacing(value)
            if ls:
                pf.line_spacing = ls
                changes.append(f"{style_name}.行距 = {ls}倍")

        elif key_lower in ("段前", "spacebefore"):
            pt = _parse_pt(value)
            if pt:
                pf.space_before = pt
                changes.append(f"{style_name}.段前 = {pt.pt}pt")

        elif key_lower in ("段后", "spaceafter"):
            pt = _parse_pt(value)
            if pt:
                pf.space_after = pt
                changes.append(f"{style_name}.段后 = {pt.pt}pt")

        elif key_lower in ("对齐", "alignment", "align"):
            if value in ALIGN_MAP:
                pf.alignment = ALIGN_MAP[value]
                changes.append(f"{style_name}.对齐 = {value}")

        elif key_lower in ("首行缩进", "firstlineindent", "indent"):
            cm = _parse_cm(value)
            if cm:
                pf.first_line_indent = cm
                changes.append(f"{style_name}.首行缩进 = {cm.cm}cm")

        elif key_lower in ("行距最小值", "linespacingatleast"):
            pt = _parse_pt(value)
            if pt:
                from docx.enum.text import WD_LINE_SPACING
                pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                pf.line_spacing = pt
                changes.append(f"{style_name}.行距 = 最小值 {pt.pt}pt")

        elif key_lower in ("行距固定值", "linespacingexact"):
            pt = _parse_pt(value)
            if pt:
                from docx.enum.text import WD_LINE_SPACING
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = pt
                changes.append(f"{style_name}.行距 = 固定值 {pt.pt}pt")

        else:
            changes.append(f"{style_name}: 未知属性 '{key}' = '{value}' (已忽略)")

    return changes


def apply_page_format(doc: Document, properties: Dict[str, str]) -> List[str]:
    """应用页面设置格式"""
    changes = []
    section = doc.sections[0]

    for key, value in properties.items():
        key_lower = key.lower()

        if key_lower in ("页边距上", "margintop", "topmargin"):
            cm = _parse_cm(value)
            if cm:
                section.top_margin = cm
                changes.append(f"页边距上 = {cm.cm}cm")

        elif key_lower in ("页边距下", "marginbottom", "bottommargin"):
            cm = _parse_cm(value)
            if cm:
                section.bottom_margin = cm
                changes.append(f"页边距下 = {cm.cm}cm")

        elif key_lower in ("页边距左", "marginleft", "leftmargin"):
            cm = _parse_cm(value)
            if cm:
                section.left_margin = cm
                changes.append(f"页边距左 = {cm.cm}cm")

        elif key_lower in ("页边距右", "marginright", "rightmargin"):
            cm = _parse_cm(value)
            if cm:
                section.right_margin = cm
                changes.append(f"页边距右 = {cm.cm}cm")

        elif key_lower in ("纸张宽度", "pagewidth"):
            cm = _parse_cm(value)
            if cm:
                section.page_width = cm
                changes.append(f"纸张宽度 = {cm.cm}cm")

        elif key_lower in ("纸张高度", "pageheight"):
            cm = _parse_cm(value)
            if cm:
                section.page_height = cm
                changes.append(f"纸张高度 = {cm.cm}cm")

        else:
            changes.append(f"页面: 未知属性 '{key}' = '{value}' (已忽略)")

    return changes


def process_document(input_path: str, output_path: str, dry_run: bool = False) -> List[str]:
    """处理文档：解析格式要求、应用、去除注释段落"""
    doc = Document(input_path)
    all_changes: List[str] = []
    paragraphs_to_remove: List[int] = []

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        req = parse_format_comment(text)
        if req:
            paragraphs_to_remove.append(idx)
            if req.target.lower() in ("页面", "page", "pagesetup"):
                changes = apply_page_format(doc, req.properties)
            else:
                changes = apply_paragraph_style_format(doc, req.target, req.properties)
            all_changes.extend(changes)

    if dry_run:
        return all_changes

    # 去除格式要求段落（从后往前删除，避免索引变化）
    for idx in reversed(paragraphs_to_remove):
        p = doc.paragraphs[idx]._element
        p.getparent().remove(p)

    # 确保输出目录存在
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    return all_changes


def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(
        description="解析模板中的 [[FORMAT: ...]] 格式要求并应用到样式"
    )
    parser.add_argument("--input", "-i", required=True, help="输入模板路径（含格式要求注释）")
    parser.add_argument("--output", "-o", required=True, help="输出去除注释后的模板路径")
    parser.add_argument("--dry-run", "-n", action="store_true", help="仅显示将要应用的变更，不实际修改")

    args = parser.parse_args(argv)

    if not Path(args.input).exists():
        print(f"错误: 输入文件不存在: {args.input}", file=sys.stderr)
        return 1

    changes = process_document(args.input, args.output, dry_run=args.dry_run)

    if args.dry_run:
        print("[DRY RUN] 以下变更将被应用:")
    else:
        print(f"已处理: {args.input} -> {args.output}")
        print("应用的格式变更:")

    if changes:
        for c in changes:
            print(f"  - {c}")
    else:
        print("  (无格式要求注释)")

    return 0


if __name__ == "__main__":
    sys.exit(main())
