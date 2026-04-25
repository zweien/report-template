#!/usr/bin/env python3
"""report-engine 模板生成器。

用法:
    python generate_template.py --mode advanced --output template.docx
    python generate_template.py --mode basic --output template.docx
    python generate_template.py --mode advanced --sections "章节1,章节2,章节3" --appendices --toc

参数:
    --mode          模板类型: basic（单章节）或 advanced（多章节+附件+条件开关）
    --output        输出路径
    --sections      章节标题列表，逗号分隔（仅 advanced 模式）
    --appendices    是否包含附件总区（仅 advanced 模式）
    --toc           是否包含目录区
    --project-name  项目名占位符的默认显示文本
"""

import argparse
from pathlib import Path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor


# ── 样式定义 ──────────────────────────────────────────────────

PARAGRAPH_STYLES = [
    # (name, options)
    ("Heading 2", {"bold": True, "color": "000000"}),
    ("Heading 3", {"bold": True, "color": "000000"}),
    ("Body Text", {}),
    ("Caption", {"italic": True}),
    ("Legend", {"italic": True}),
    ("Figure Paragraph", {}),
    ("List Bullet", {}),
    ("List Number", {}),
    ("Note", {"italic": True}),
    ("Quote", {"italic": True}),
    ("Checklist", {}),
    ("CodeBlock", {"font_name": "Courier New"}),
]

TABLE_STYLES = ["ResearchTable", "AppendixTable"]

# ── 占位符定义 ──────────────────────────────────────────────────

SCALAR_PLACEHOLDERS = [
    ("PROJECT_NAME", "项目名称"),
    ("APPLICANT_ORG", "申请单位"),
    ("PROJECT_LEADER", "项目负责人"),
    ("PROJECT_PERIOD", "项目周期"),
]

# ── 默认章节 ──────────────────────────────────────────────────

DEFAULT_SECTIONS = [
    ("研究内容与技术路线", "RESEARCH_CONTENT", 1),
    ("研究基础与条件保障", "RESEARCH_BASIS", 2),
    ("实施计划与进度安排", "IMPLEMENTATION_PLAN", 3),
]


def add_styles(doc: Document):
    """添加所有必需样式。"""
    for name, opts in PARAGRAPH_STYLES:
        try:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            style = doc.styles[name]
        if opts.get("bold"):
            style.font.bold = True
        if opts.get("italic"):
            style.font.italic = True
        if opts.get("font_name"):
            style.font.name = opts["font_name"]
        if opts.get("color"):
            style.font.color.rgb = RGBColor.from_string(opts["color"])

    for name in TABLE_STYLES:
        try:
            doc.styles.add_style(name, WD_STYLE_TYPE.TABLE)
        except ValueError:
            pass


def add_cover(doc: Document):
    """添加封面页。"""
    doc.add_paragraph("{{PROJECT_NAME}}", style="Heading 2")
    doc.add_paragraph("申请单位：{{APPLICANT_ORG}}")
    doc.add_paragraph("项目负责人：{{PROJECT_LEADER}}")
    doc.add_paragraph("项目周期：{{PROJECT_PERIOD}}")
    doc.add_page_break()


def add_toc(doc: Document):
    """添加目录区。"""
    doc.add_paragraph("{%p if ENABLE_TOC %}")
    doc.add_paragraph("{{p TOC_SUBDOC }}")
    doc.add_paragraph("{%p endif %}")
    doc.add_page_break()


def add_section(doc: Document, title: str, prefix: str, number: int):
    """添加一个带条件开关的章节。"""
    chinese_nums = "一二三四五六七八九十"
    num_str = chinese_nums[number - 1] if number <= len(chinese_nums) else str(number)
    doc.add_paragraph(f"{{%p if ENABLE_{prefix} %}}")
    doc.add_paragraph(f"{num_str}、{title}")
    doc.add_paragraph(f"{{{{p {prefix}_SUBDOC }}}}")
    doc.add_paragraph("{%p endif %}")


def add_appendices(doc: Document):
    """添加附件总区。"""
    doc.add_paragraph("{%p if ENABLE_APPENDICES %}")
    doc.add_paragraph("附件")
    doc.add_paragraph("{{p APPENDICES_SUBDOC }}")
    doc.add_paragraph("{%p endif %}")


def build_basic(output_path: str):
    """生成基础版模板（单章节）。"""
    doc = Document()
    add_styles(doc)
    add_cover(doc)

    doc.add_paragraph("研究内容")
    doc.add_paragraph("{{p RESEARCH_CONTENT_SUBDOC }}")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"基础版模板已生成: {output_path}")


def build_advanced(output_path: str, sections=None, include_appendices=True,
                   include_toc=True):
    """生成进阶版模板（多章节+附件+条件开关）。"""
    if sections is None:
        section_defs = DEFAULT_SECTIONS
    else:
        section_defs = []
        for i, title in enumerate(sections, 1):
            # 从标题生成 prefix：取前几个字的大写英文
            prefix = f"SECTION_{i}"
            section_defs.append((title, prefix, i))

    doc = Document()
    add_styles(doc)
    add_cover(doc)

    if include_toc:
        add_toc(doc)

    for title, prefix, number in section_defs:
        add_section(doc, title, prefix, number)

    if include_appendices:
        add_appendices(doc)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"进阶版模板已生成: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="report-engine 模板生成器")
    parser.add_argument("--mode", choices=["basic", "advanced"], default="advanced",
                        help="模板类型")
    parser.add_argument("--output", required=True, help="输出路径")
    parser.add_argument("--sections", help="章节标题，逗号分隔")
    parser.add_argument("--appendices", action="store_true", default=True,
                        help="包含附件总区")
    parser.add_argument("--no-appendices", action="store_true",
                        help="不包含附件总区")
    parser.add_argument("--toc", action="store_true", help="包含目录区")
    args = parser.parse_args()

    if args.mode == "basic":
        build_basic(args.output)
    else:
        sections = args.sections.split(",") if args.sections else None
        build_advanced(
            args.output,
            sections=sections,
            include_appendices=not args.no_appendices,
            include_toc=args.toc,
        )


if __name__ == "__main__":
    main()
