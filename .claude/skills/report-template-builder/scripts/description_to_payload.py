#!/usr/bin/env python3
"""根据文字描述生成 report-engine 模板和 payload。

用法:
    python description_to_payload.py --description "项目申报书，包含研究内容、研究基础、实施计划、附件" \
      --output-template templates/generated.docx \
      --output-payload data/examples/generated.json

    python description_to_payload.py --file description.txt \
      --output-template templates/generated.docx \
      --output-payload data/examples/generated.json

描述格式:
    自由文本，程序会尝试识别：
    - 章节标题（中文数字序号：一、二、三 或 第X章）
    - 关键词：目录、附件、附录、封面
    - 分隔符：逗号、顿号、换行
"""

import argparse
import json
import re
from pathlib import Path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE


# ── 样式定义 ──────────────────────────────────────────────────

REQUIRED_PARAGRAPH_STYLES = [
    ("Heading 2", {"bold": True}),
    ("Heading 3", {"bold": True}),
    ("Body Text", {}),
    ("Caption", {"italic": True}),
    ("Legend", {"italic": True}),
    ("Figure Paragraph", {}),
    ("List Bullet", {}),
    ("List Number", {}),
    ("Note", {"italic": True}),
    ("Quote", {"italic": True}),
    ("Checklist", {}),
    ("CodeBlock", {"font_name": "Courier New"}),
]

REQUIRED_TABLE_STYLES = ["ResearchTable", "AppendixTable"]


def parse_description(text: str) -> dict:
    """解析描述文本，提取结构信息。"""
    text = text.strip()

    result = {
        "title": "",
        "sections": [],
        "has_toc": False,
        "has_appendices": False,
        "extra_info": [],
    }

    # 检测目录
    if re.search(r"目录|TOC|table.of.contents", text, re.IGNORECASE):
        result["has_toc"] = True

    # 检测附件
    if re.search(r"附件|附录|appendix|appendices", text, re.IGNORECASE):
        result["has_appendices"] = True

    # 提取标题（第一个出现的项目名/报告名）
    title_match = re.search(r"[《「]([^》」]+)[》」]", text)
    if title_match:
        result["title"] = title_match.group(1)

    # 提取章节（多种模式）
    sections = []

    # 模式1：中文数字序号 一、二、三、...
    chinese_pattern = r"[一二三四五六七八九十]+[、．.]\s*([^\n,，。]+)"
    for m in re.finditer(chinese_pattern, text):
        sections.append(m.group(1).strip())

    # 模式2：第X章 / 第X节
    chapter_pattern = r"第[一二三四五六七八九十\d]+[章节]\s*[：:]?\s*([^\n,，。]+)"
    for m in re.finditer(chapter_pattern, text):
        sections.append(m.group(1).strip())

    # 模式3：数字序号 1. 2. 3.
    num_pattern = r"(\d+)[.、）)]\s*([^\n,，。]+)"
    for m in re.finditer(num_pattern, text):
        sections.append(m.group(2).strip())

    # 模式4：逗号/顿号分隔的列表（如果没有找到带序号的章节）
    if not sections:
        # 去掉已识别的关键词和常见填充词
        cleaned = re.sub(r"目录|附件|附录|封面|项目申报书|报告|开题|结题|中期", "", text)
        cleaned = re.sub(r"包含|包括|需要|设有|以及|和|与|，|。", ",", cleaned)
        # 去掉冒号前的引导文字
        cleaned = re.sub(r"^[^：:]+[：:]\s*", "", cleaned)
        parts = re.split(r"[,，、；;\n]+", cleaned)
        # 过滤掉太短的、明显是填充词的
        stopwords = {"包含", "包括", "需要", "设有", "以及", "和", "与", "等"}
        sections = [p.strip() for p in parts if len(p.strip()) >= 2 and p.strip() not in stopwords]

    # 去重并保持顺序
    seen = set()
    for s in sections:
        if s not in seen and s not in ("目录", "附件", "附录"):
            seen.add(s)
            result["sections"].append(s)

    return result


def generate_template(parsed: dict, output_path: str):
    """根据解析结果生成模板。"""
    doc = Document()

    # 添加样式
    for name, opts in REQUIRED_PARAGRAPH_STYLES:
        try:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            if opts.get("bold"):
                style.font.bold = True
            if opts.get("italic"):
                style.font.italic = True
            if opts.get("font_name"):
                style.font.name = opts["font_name"]
        except ValueError:
            pass

    for name in REQUIRED_TABLE_STYLES:
        try:
            doc.styles.add_style(name, WD_STYLE_TYPE.TABLE)
        except ValueError:
            pass

    # 封面
    doc.add_paragraph("{{PROJECT_NAME}}", style="Heading 2")
    doc.add_paragraph("申请单位：{{APPLICANT_ORG}}")
    doc.add_paragraph("项目负责人：{{PROJECT_LEADER}}")
    doc.add_paragraph("项目周期：{{PROJECT_PERIOD}}")
    doc.add_page_break()

    # 目录
    if parsed["has_toc"]:
        doc.add_paragraph("{%p if ENABLE_TOC %}")
        doc.add_paragraph("{{p TOC_SUBDOC }}")
        doc.add_paragraph("{%p endif %}")
        doc.add_page_break()

    # 章节
    chinese_nums = "一二三四五六七八九十"
    section_defs = []
    for i, title in enumerate(parsed["sections"][:10]):
        prefix = f"SECTION_{i + 1}"
        num = chinese_nums[i] if i < len(chinese_nums) else str(i + 1)
        section_defs.append((title, prefix, num))
        doc.add_paragraph(f"{{%p if ENABLE_{prefix} %}}")
        doc.add_paragraph(f"{num}、{title}")
        doc.add_paragraph(f"{{{{p {prefix}_SUBDOC }}}}")
        doc.add_paragraph("{%p endif %}")

    # 附件
    if parsed["has_appendices"]:
        doc.add_paragraph("{%p if ENABLE_APPENDICES %}")
        doc.add_paragraph("附件")
        doc.add_paragraph("{{p APPENDICES_SUBDOC }}")
        doc.add_paragraph("{%p endif %}")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"模板已生成: {output_path}")
    return section_defs


def generate_payload(parsed: dict, section_defs: list) -> dict:
    """根据解析结果生成 payload。"""
    title = parsed["title"] or "自动生成报告"

    context = {
        "PROJECT_NAME": title,
        "APPLICANT_ORG": "（请填写）",
        "PROJECT_LEADER": "（请填写）",
        "PROJECT_PERIOD": "（请填写）",
    }

    sections = []
    order = 0

    if parsed["has_toc"]:
        sections.append({
            "id": "toc",
            "placeholder": "TOC_SUBDOC",
            "flag_name": "ENABLE_TOC",
            "enabled": True,
            "blocks": [{"type": "toc_placeholder", "title": "目 录"}],
            "order": order,
        })
        order += 1

    for title, prefix, _ in section_defs:
        sections.append({
            "id": prefix.lower(),
            "placeholder": f"{prefix}_SUBDOC",
            "flag_name": f"ENABLE_{prefix}",
            "enabled": True,
            "blocks": [
                {"type": "heading", "text": "（请填写小节标题）", "level": 2},
                {"type": "paragraph", "text": "（请填写正文内容）"},
            ],
            "order": order,
        })
        order += 1

    attachments_bundle = None
    if parsed["has_appendices"]:
        attachments_bundle = {
            "enabled": True,
            "placeholder": "APPENDICES_SUBDOC",
            "flag_name": "ENABLE_APPENDICES",
            "page_break_between_attachments": True,
            "include_attachment_title": True,
        }

    return {
        "context": context,
        "sections": sections,
        "attachments": [],
        "attachments_bundle": attachments_bundle,
        "style_map": {
            "table": "ResearchTable",
            "figure_paragraph": "Figure Paragraph",
            "legend": "Legend",
        },
    }


def main():
    parser = argparse.ArgumentParser(description="根据描述生成 report-engine 模板和 payload")
    parser.add_argument("--description", help="描述文本")
    parser.add_argument("--file", help="包含描述的文本文件")
    parser.add_argument("--output-template", required=True, help="输出模板路径")
    parser.add_argument("--output-payload", required=True, help="输出 payload 路径")
    args = parser.parse_args()

    if args.description:
        text = args.description
    elif args.file:
        text = Path(args.file).read_text(encoding="utf-8")
    else:
        parser.error("需要 --description 或 --file 参数")

    print(f"解析描述: {text[:100]}...")
    parsed = parse_description(text)

    print(f"  标题: {parsed['title'] or '（未识别）'}")
    print(f"  章节: {parsed['sections']}")
    print(f"  目录: {'是' if parsed['has_toc'] else '否'}")
    print(f"  附件: {'是' if parsed['has_appendices'] else '否'}")

    section_defs = generate_template(parsed, args.output_template)

    payload = generate_payload(parsed, section_defs)
    Path(args.output_payload).parent.mkdir(parents=True, exist_ok=True)
    with open(args.output_payload, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)
    print(f"Payload 已生成: {args.output_payload}")

    print("\n下一步：")
    print(f"  1. 编辑 {args.output_payload} 填写具体内容")
    print(f"  2. 运行 report-engine check-template --template {args.output_template} --payload {args.output_payload}")
    print(f"  3. 运行 report-engine render --template {args.output_template} --payload {args.output_payload} --output output/result.docx")


if __name__ == "__main__":
    main()
