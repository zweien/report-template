#!/usr/bin/env python3
"""分析现有 .docx 文件，提取结构信息，生成 report-engine 兼容的模板和 payload。

用法:
    python analyze_docx.py --input reference.docx --output-template templates/generated.docx --output-payload data/examples/generated.json

功能:
    1. 读取输入 docx 的标题层级、段落样式、表格结构
    2. 映射到 report-engine 的 section 结构
    3. 生成包含必需样式和占位符的模板
    4. 生成匹配的 payload（含占位 block）
"""

import argparse
import json
import re
from pathlib import Path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE


# ── 样式定义 ──────────────────────────────────────────────────

REQUIRED_PARAGRAPH_STYLES = [
    ("Heading 2", {"bold": True}),
    ("Heading 3", {"bold": True}),
    ("Body Text", {}),
    ("Caption", {"italic": True}),
    ("Legend", {"italic": True}),
    ("Figure Paragraph", {}),
    ("List Bullet", {}),
    ("List Number", {}),
    ("Note", {"italic": True}),
    ("Quote", {"italic": True}),
    ("Checklist", {}),
    ("CodeBlock", {"font_name": "Courier New"}),
]

REQUIRED_TABLE_STYLES = ["ResearchTable", "AppendixTable"]


def extract_structure(input_path: str) -> dict:
    """从 docx 中提取文档结构。"""
    doc = Document(input_path)

    headings = []
    tables_count = 0
    paragraphs_count = 0
    styles_used = set()

    for para in doc.paragraphs:
        paragraphs_count += 1
        style_name = para.style.name if para.style else "Normal"
        styles_used.add(style_name)

        # 识别标题
        if style_name.startswith("Heading") and para.text.strip():
            level = int(style_name.replace("Heading ", "")) if style_name[-1].isdigit() else 2
            headings.append({
                "text": para.text.strip(),
                "level": level,
                "style": style_name,
            })

    for table in doc.tables:
        tables_count += 1

    # 识别可能的章节结构（基于 Heading 1/2）
    sections = []
    for h in headings:
        if h["level"] <= 2:
            sections.append(h["text"])

    return {
        "headings": headings,
        "tables_count": tables_count,
        "paragraphs_count": paragraphs_count,
        "styles_used": sorted(styles_used),
        "sections": sections,
        "has_cover": any("项目" in h["text"] or "申报" in h["text"]
                         for h in headings if h["level"] == 1),
    }


def generate_template_from_structure(structure: dict, output_path: str,
                                      include_toc: bool = True,
                                      include_appendices: bool = True):
    """根据提取的结构生成 report-engine 模板。"""
    doc = Document()

    # 添加必需样式
    for name, opts in REQUIRED_PARAGRAPH_STYLES:
        try:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            if opts.get("bold"):
                style.font.bold = True
            if opts.get("italic"):
                style.font.italic = True
            if opts.get("font_name"):
                style.font.name = opts["font_name"]
        except ValueError:
            pass

    for name in REQUIRED_TABLE_STYLES:
        try:
            doc.styles.add_style(name, WD_STYLE_TYPE.TABLE)
        except ValueError:
            pass

    # 封面
    doc.add_paragraph("{{PROJECT_NAME}}", style="Heading 2")
    doc.add_paragraph("申请单位：{{APPLICANT_ORG}}")
    doc.add_paragraph("项目负责人：{{PROJECT_LEADER}}")
    doc.add_paragraph("项目周期：{{PROJECT_PERIOD}}")
    doc.add_page_break()

    # 目录
    if include_toc:
        doc.add_paragraph("{%p if ENABLE_TOC %}")
        doc.add_paragraph("{{p TOC_SUBDOC }}")
        doc.add_paragraph("{%p endif %}")
        doc.add_page_break()

    # 章节
    chinese_nums = "一二三四五六七八九十"
    section_defs = []
    for i, title in enumerate(structure["sections"][:10]):  # 最多 10 个章节
        prefix = f"SECTION_{i + 1}"
        num = chinese_nums[i] if i < len(chinese_nums) else str(i + 1)
        section_defs.append((title, prefix, num))
        doc.add_paragraph(f"{{%p if ENABLE_{prefix} %}}")
        doc.add_paragraph(f"{num}、{title}")
        doc.add_paragraph(f"{{{{p {prefix}_SUBDOC }}}}")
        doc.add_paragraph("{%p endif %}")

    # 附件
    if include_appendices:
        doc.add_paragraph("{%p if ENABLE_APPENDICES %}")
        doc.add_paragraph("附件")
        doc.add_paragraph("{{p APPENDICES_SUBDOC }}")
        doc.add_paragraph("{%p endif %}")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"模板已生成: {output_path}")
    return section_defs


def generate_payload_from_structure(structure: dict, section_defs: list,
                                     include_toc: bool = True,
                                     include_appendices: bool = True) -> dict:
    """根据提取的结构生成 payload。"""
    context = {
        "PROJECT_NAME": "自动生成报告",
        "APPLICANT_ORG": "（请填写）",
        "PROJECT_LEADER": "（请填写）",
        "PROJECT_PERIOD": "（请填写）",
    }

    sections = []
    order = 0

    # 目录
    if include_toc:
        sections.append({
            "id": "toc",
            "placeholder": "TOC_SUBDOC",
            "flag_name": "ENABLE_TOC",
            "enabled": True,
            "blocks": [{"type": "toc_placeholder", "title": "目 录"}],
            "order": order,
        })
        order += 1

    # 章节
    for title, prefix, _ in section_defs:
        sections.append({
            "id": prefix.lower(),
            "placeholder": f"{prefix}_SUBDOC",
            "flag_name": f"ENABLE_{prefix}",
            "enabled": True,
            "subdoc_title": title,
            "subdoc_title_level": 2,
            "blocks": [
                {"type": "heading", "text": "（请填写小节标题）", "level": 2},
                {"type": "paragraph", "text": "（请填写正文内容）"},
            ],
            "order": order,
        })
        order += 1

    # 附件
    attachments = []
    attachments_bundle = None
    if include_appendices:
        attachments_bundle = {
            "enabled": True,
            "placeholder": "APPENDICES_SUBDOC",
            "flag_name": "ENABLE_APPENDICES",
            "page_break_between_attachments": True,
            "include_attachment_title": True,
        }

    payload = {
        "context": context,
        "sections": sections,
        "attachments": attachments,
        "attachments_bundle": attachments_bundle,
        "style_map": {
            "table": "ResearchTable",
            "figure_paragraph": "Figure Paragraph",
            "legend": "Legend",
        },
    }

    return payload


def main():
    parser = argparse.ArgumentParser(description="分析 docx 并生成 report-engine 模板和 payload")
    parser.add_argument("--input", required=True, help="输入的 docx 文件路径")
    parser.add_argument("--output-template", required=True, help="输出模板路径")
    parser.add_argument("--output-payload", required=True, help="输出 payload 路径")
    parser.add_argument("--no-toc", action="store_true", help="不包含目录")
    parser.add_argument("--no-appendices", action="store_true", help="不包含附件区")
    args = parser.parse_args()

    print(f"分析输入文件: {args.input}")
    structure = extract_structure(args.input)

    print(f"  标题数: {len(structure['headings'])}")
    print(f"  表格数: {structure['tables_count']}")
    print(f"  段落数: {structure['paragraphs_count']}")
    print(f"  识别到的章节: {structure['sections']}")

    section_defs = generate_template_from_structure(
        structure, args.output_template,
        include_toc=not args.no_toc,
        include_appendices=not args.no_appendices,
    )

    payload = generate_payload_from_structure(
        structure, section_defs,
        include_toc=not args.no_toc,
        include_appendices=not args.no_appendices,
    )

    Path(args.output_payload).parent.mkdir(parents=True, exist_ok=True)
    with open(args.output_payload, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)
    print(f"Payload 已生成: {args.output_payload}")

    print("\n下一步：")
    print(f"  1. 编辑 {args.output_payload} 填写具体内容")
    print(f"  2. 运行 report-engine check-template --template {args.output_template} --payload {args.output_payload}")
    print(f"  3. 运行 report-engine render --template {args.output_template} --payload {args.output_payload} --output output/result.docx")


if __name__ == "__main__":
    main()
