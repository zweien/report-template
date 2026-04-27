from __future__ import annotations

import pytest
from docx import Document

from report_engine.prompt_parser import (
    PROMPT_PREFIX,
    PROMPT_SUFFIX,
    _parse_prompt_text,
    extract_prompts,
    filter_prompt_paragraphs,
)


def test_parse_full_prompt():
    text = "[[PROMPT: 立项依据: 请从国内外研究现状撰写 | mode=auto]]"
    result = _parse_prompt_text(text)
    assert result is not None
    assert result["target"] == "立项依据"
    assert result["prompt"] == "请从国内外研究现状撰写"
    assert result["mode"] == "auto"
    assert result["level"] == "section"


def test_parse_prompt_without_mode():
    text = "[[PROMPT: 研究目标: 列出3-5个具体目标]]"
    result = _parse_prompt_text(text)
    assert result is not None
    assert result["target"] == "研究目标"
    assert result["prompt"] == "列出3-5个具体目标"
    assert result["mode"] == "auto"
    assert result["level"] == "section"


def test_parse_paragraph_level_prompt():
    text = "[[PROMPT: 研究内容.技术路线: 描述技术路线 | mode=interactive]]"
    result = _parse_prompt_text(text)
    assert result is not None
    assert result["target"] == "研究内容.技术路线"
    assert result["prompt"] == "描述技术路线"
    assert result["mode"] == "interactive"
    assert result["level"] == "paragraph"


def test_parse_invalid_mode_defaults_to_auto():
    text = "[[PROMPT: 研究基础: 描述研究基础 | mode=invalid]]"
    result = _parse_prompt_text(text)
    assert result is not None
    assert result["target"] == "研究基础"
    assert result["prompt"] == "描述研究基础"
    assert result["mode"] == "auto"
    assert result["level"] == "section"


def test_parse_multiline_prompt():
    text = "[[PROMPT: 研究内容:\n请详细描述研究内容\n包括多个方面\n | mode=interactive]]"
    result = _parse_prompt_text(text)
    assert result is not None
    assert result["target"] == "研究内容"
    assert result["prompt"] == "请详细描述研究内容\n包括多个方面"
    assert result["mode"] == "interactive"
    assert result["level"] == "section"


def test_parse_empty_target_returns_none():
    text = "[[PROMPT: : 只有提示没有目标]]"
    result = _parse_prompt_text(text)
    assert result is None


def test_parse_non_prompt_returns_none():
    assert _parse_prompt_text("普通文本") is None
    assert _parse_prompt_text("[[PROMPT: 缺少后缀") is None
    assert _parse_prompt_text("缺少前缀: 目标: 提示]]") is None
    assert _parse_prompt_text("") is None


def test_extract_prompts():
    doc = Document()
    doc.add_paragraph("[[PROMPT: 目标1: 提示1 | mode=auto]]")
    doc.add_paragraph("普通段落")
    doc.add_paragraph("[[PROMPT: 目标2: 提示2 | mode=interactive]]")

    prompts = extract_prompts(doc)
    assert len(prompts) == 2
    assert prompts[0]["target"] == "目标1"
    assert prompts[1]["target"] == "目标2"


def test_filter_prompt_paragraphs():
    doc = Document()
    doc.add_paragraph("[[PROMPT: 目标1: 提示1 | mode=auto]]")
    doc.add_paragraph("普通段落")
    doc.add_paragraph("[[PROMPT: 目标2: 提示2 | mode=interactive]]")
    doc.add_paragraph("另一个普通段落")

    removed = filter_prompt_paragraphs(doc)
    assert removed == 2

    remaining_texts = [p.text for p in doc.paragraphs]
    assert remaining_texts == ["普通段落", "另一个普通段落"]


def test_filter_prompt_paragraphs_no_prompts():
    doc = Document()
    doc.add_paragraph("普通段落1")
    doc.add_paragraph("普通段落2")

    removed = filter_prompt_paragraphs(doc)
    assert removed == 0
    assert len(doc.paragraphs) == 2
