from __future__ import annotations

import json
import sys
from pathlib import Path
from typing import Any, Dict

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docxtpl import DocxTemplate

from report_engine.blocks import create_default_registry

REPO_ROOT = Path(__file__).resolve().parent.parent
SRC_DIR = REPO_ROOT / "src"
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))


@pytest.fixture
def registry():
    return create_default_registry()


@pytest.fixture
def project_root() -> Path:
    return REPO_ROOT


@pytest.fixture
def minimal_template(tmp_path: Path) -> str:
    doc = Document()

    paragraph_styles = [
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Heading 4",
        "Heading 5",
        "Body Text",
        "TableCaption",
        "FigureCaption",
        "Legend",
        "Figure Paragraph",
        "List Bullet",
        "List Number",
        "Note",
        "Quote",
        "Checklist",
        "CodeBlock",
    ]
    for name in paragraph_styles:
        try:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            pass

    for table_style_name in ["ResearchTable", "AppendixTable"]:
        try:
            doc.styles.add_style(table_style_name, WD_STYLE_TYPE.TABLE)
        except ValueError:
            pass

    doc.add_paragraph("{{PROJECT_NAME}}")
    doc.add_paragraph("{{APPLICANT_ORG}}")
    doc.add_paragraph("{%p if ENABLE_RESEARCH_CONTENT %}")
    doc.add_paragraph("二、研究内容与技术路线")
    doc.add_paragraph("{{p RESEARCH_CONTENT_SUBDOC }}")
    doc.add_paragraph("{%p endif %}")
    doc.add_paragraph("{%p if ENABLE_RESEARCH_BASIS %}")
    doc.add_paragraph("三、研究基础")
    doc.add_paragraph("{{p RESEARCH_BASIS_SUBDOC }}")
    doc.add_paragraph("{%p endif %}")
    doc.add_paragraph("{%p if ENABLE_APPENDICES %}")
    doc.add_paragraph("附件")
    doc.add_paragraph("{{p APPENDICES_SUBDOC }}")
    doc.add_paragraph("{%p endif %}")

    path = tmp_path / "test_template.docx"
    doc.save(path)
    return str(path)


@pytest.fixture
def tpl(minimal_template: str) -> DocxTemplate:
    return DocxTemplate(minimal_template)


@pytest.fixture
def advanced_payload() -> Dict[str, Any]:
    return {
        "context": {
            "PROJECT_NAME": "测试项目",
            "APPLICANT_ORG": "测试单位",
        },
        "sections": [
            {
                "id": "research_content",
                "placeholder": "RESEARCH_CONTENT_SUBDOC",
                "flag_name": "ENABLE_RESEARCH_CONTENT",
                "enabled": True,
                "blocks": [
                    {"type": "heading", "text": "研究目标", "level": 2},
                    {"type": "paragraph", "text": "这是测试正文。"},
                ],
            },
            {
                "id": "research_basis",
                "placeholder": "RESEARCH_BASIS_SUBDOC",
                "flag_name": "ENABLE_RESEARCH_BASIS",
                "enabled": False,
                "blocks": [
                    {"type": "paragraph", "text": "不会输出。"},
                ],
            },
        ],
        "attachments": [
            {
                "id": "appendix_1",
                "placeholder": "APPENDIX_1_SUBDOC",
                "flag_name": "ENABLE_APPENDIX_1",
                "enabled": True,
                "title": "附件1",
                "blocks": [
                    {"type": "paragraph", "text": "附件内容。"},
                ],
            }
        ],
        "attachments_bundle": {
            "enabled": True,
            "placeholder": "APPENDICES_SUBDOC",
            "flag_name": "ENABLE_APPENDICES",
            "page_break_between_attachments": True,
            "include_attachment_title": True,
        },
        "style_map": {},
    }


@pytest.fixture
def payload_path(tmp_path: Path, advanced_payload: Dict[str, Any]) -> str:
    path = tmp_path / "payload.json"
    path.write_text(json.dumps(advanced_payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return str(path)
