from report_engine.schema import Payload
from report_engine.template_checker import check_template_contract


def test_template_checker_accepts_minimal_template(minimal_template, advanced_payload):
    payload = Payload.model_validate(advanced_payload)
    result = check_template_contract(minimal_template, payload)
    assert result.ok is True
    assert result.missing_placeholders == []


def test_template_checker_reports_missing_section_placeholder(tmp_path, advanced_payload):
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()
    for name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5", "Body Text", "Caption", "Legend", "Figure Paragraph", "List Bullet", "List Number"]:
        try:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            pass
    try:
        doc.styles.add_style("ResearchTable", WD_STYLE_TYPE.TABLE)
    except ValueError:
        pass

    doc.add_paragraph("{{PROJECT_NAME}}")
    doc.add_paragraph("{%p if ENABLE_APPENDICES %}")
    doc.add_paragraph("{{p APPENDICES_SUBDOC }}")
    doc.add_paragraph("{%p endif %}")
    path = tmp_path / "missing_section.docx"
    doc.save(path)

    payload = Payload.model_validate(advanced_payload)
    result = check_template_contract(str(path), payload)
    assert result.ok is False
    assert "RESEARCH_CONTENT_SUBDOC" in result.missing_placeholders


def test_template_checker_allows_bundle_without_individual_attachment_placeholders(minimal_template, advanced_payload):
    payload = Payload.model_validate(advanced_payload)
    result = check_template_contract(minimal_template, payload)
    assert result.ok is True
    assert any("bundled appendix slot" in note for note in result.notes)


def test_template_checker_warns_missing_context_vars(minimal_template, advanced_payload):
    # Remove a context var that the template uses
    payload_data = dict(advanced_payload)
    payload_data["context"] = {}
    payload = Payload.model_validate(payload_data)
    result = check_template_contract(minimal_template, payload)
    assert any("PROJECT_NAME" in w for w in result.warnings)
    assert any("APPLICANT_ORG" in w for w in result.warnings)
