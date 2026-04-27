from docx import Document

from report_engine.renderer import render_report


def test_render_report_outputs_docx(minimal_template, advanced_payload, tmp_path):
    output_path = tmp_path / "result.docx"
    warnings = render_report(minimal_template, str(output_path), advanced_payload)
    assert output_path.exists()
    assert output_path.stat().st_size > 0
    assert isinstance(warnings, list)

    doc = Document(str(output_path))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # context 变量已渲染
    assert "测试项目" in full_text
    assert "测试单位" in full_text

    # enabled section 的内容存在
    assert "研究目标" in full_text
    assert "这是测试正文。" in full_text

    # disabled section 的内容不应存在
    assert "不会输出。" not in full_text

    # attachment 内容通过 bundle 聚合后存在
    assert "附件内容。" in full_text


def test_render_report_can_skip_template_checks(minimal_template, advanced_payload, tmp_path):
    broken_payload = dict(advanced_payload)
    broken_payload["attachments_bundle"] = None
    output_path = tmp_path / "result_no_checks.docx"
    warnings = render_report(
        minimal_template,
        str(output_path),
        broken_payload,
        check_template=False,
    )
    assert output_path.exists()
    assert isinstance(warnings, list)

    doc = Document(str(output_path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "测试项目" in full_text


def test_render_report_filters_prompt_paragraphs(minimal_template, advanced_payload, tmp_path):
    """渲染时应自动删除 [[PROMPT: ...]] 段落。"""
    from docx import Document

    doc = Document(minimal_template)
    # Insert a PROMPT paragraph before the first conditional block
    prompt_para = doc.add_paragraph("[[PROMPT: 研究内容: 请撰写研究内容 | mode=auto]]")
    body = doc.element.body
    body.insert(2, prompt_para._element)
    doc.save(minimal_template)

    output_path = tmp_path / "result_no_prompt.docx"
    render_report(minimal_template, str(output_path), advanced_payload)

    doc = Document(str(output_path))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # PROMPT text should NOT appear in output
    assert "[[PROMPT:" not in full_text
    assert "请撰写研究内容" not in full_text
    # Normal content should still be present
    assert "测试项目" in full_text
