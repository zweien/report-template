import pytest

from unittest.mock import patch

import pytest

from report_engine.blocks import BlockRenderError, BlockRegistry, create_default_registry


@pytest.fixture
def subdoc(tpl):
    return tpl.new_subdoc()


@pytest.fixture
def style_map():
    return {
        "heading_1": "Heading 1",
        "heading_2": "Heading 2",
        "heading_3": "Heading 3",
        "heading_4": "Heading 4",
        "heading_5": "Heading 5",
        "body": "Body Text",
        "table_caption": "TableCaption",
        "figure_caption": "FigureCaption",
        "legend": "Legend",
        "figure_paragraph": "Figure Paragraph",
        "table": "ResearchTable",
        "bullet_list": "List Bullet",
        "numbered_list": "List Number",
    }


def test_unknown_block_type_raises(subdoc, style_map):
    registry = BlockRegistry()
    with pytest.raises(BlockRenderError, match="Unsupported block type"):
        registry.render(subdoc, {"type": "unknown"}, style_map)


def test_heading_block(subdoc, style_map):
    registry = create_default_registry()
    registry.render(subdoc, {"type": "heading", "text": "标题", "level": 2}, style_map)
    assert subdoc.paragraphs[0].text == "标题"
    assert subdoc.paragraphs[0].style.name == "Heading 2"


def test_heading_block_levels(subdoc, style_map):
    registry = create_default_registry()
    for level in [1, 2, 3, 4, 5]:
        registry.render(subdoc, {"type": "heading", "text": f"标题{level}", "level": level}, style_map)
    for i, level in enumerate([1, 2, 3, 4, 5], start=0):
        assert subdoc.paragraphs[i].style.name == f"Heading {level}"


def test_heading_block_clamps_level(subdoc, style_map):
    registry = create_default_registry()
    registry.render(subdoc, {"type": "heading", "text": "太低", "level": 0}, style_map)
    assert subdoc.paragraphs[0].style.name == "Heading 1"
    registry.render(subdoc, {"type": "heading", "text": "太高", "level": 99}, style_map)
    assert subdoc.paragraphs[1].style.name == "Heading 5"


def test_paragraph_block(subdoc, style_map):
    registry = create_default_registry()
    registry.render(subdoc, {"type": "paragraph", "text": "正文"}, style_map)
    assert subdoc.paragraphs[0].text == "正文"


def test_bullet_list_block(subdoc, style_map):
    registry = create_default_registry()
    registry.render(subdoc, {"type": "bullet_list", "items": ["A", "B"]}, style_map)
    assert len(subdoc.paragraphs) == 2
    assert subdoc.paragraphs[0].text == "A"


def test_numbered_list_block(subdoc, style_map):
    registry = create_default_registry()
    registry.render(subdoc, {"type": "numbered_list", "items": ["1", "2"]}, style_map)
    assert len(subdoc.paragraphs) == 2


def test_table_block(subdoc, style_map):
    registry = create_default_registry()
    registry.render(
        subdoc,
        {"type": "table", "title": "表1", "headers": ["A"], "rows": [["1"]]},
        style_map,
    )
    assert len(subdoc.tables) == 1
    assert subdoc.tables[0].rows[0].cells[0].text == "A"
    assert subdoc.tables[0].rows[1].cells[0].text == "1"


def test_image_block_missing_file(subdoc, style_map):
    registry = create_default_registry()
    registry.render(subdoc, {"type": "image", "path": "missing.png"}, style_map)
    assert any("图片缺失" in p.text for p in subdoc.paragraphs)

def test_rich_paragraph_block(subdoc, style_map, registry):
    block = {
        "type": "rich_paragraph",
        "segments": [
            {"text": "普通文本 "},
            {"text": "粗体", "bold": True},
            {"text": " 斜体", "italic": True},
            {"text": "H", "sub": True},
            {"text": "2O"},
        ],
    }
    registry.render(subdoc, block, style_map)
    assert len(subdoc.paragraphs) == 1
    p = subdoc.paragraphs[0]
    assert len(p.runs) == 5
    assert p.runs[0].text == "普通文本 "
    assert p.runs[1].bold is True
    assert p.runs[2].italic is True

def test_note_block(subdoc, style_map, registry):
    block = {"type": "note", "text": "本表数据为示意数据。"}
    registry.render(subdoc, block, style_map)
    assert len(subdoc.paragraphs) == 1
    p = subdoc.paragraphs[0]
    assert p.runs[0].text == "注："
    assert p.runs[0].bold is True
    assert p.runs[1].text == "本表数据为示意数据。"

def test_quote_block_with_source(subdoc, style_map, registry):
    block = {
        "type": "quote",
        "text": "教育是国之大计、党之大计。",
        "source": "《中国教育现代化2035》",
    }
    registry.render(subdoc, block, style_map)
    assert len(subdoc.paragraphs) == 2
    assert "教育是国之大计" in subdoc.paragraphs[0].text
    assert "中国教育现代化" in subdoc.paragraphs[1].text


def test_quote_block_without_source(subdoc, style_map, registry):
    block = {"type": "quote", "text": "引用文本。"}
    registry.render(subdoc, block, style_map)
    assert len(subdoc.paragraphs) == 1

def test_two_images_row_block(subdoc, style_map, registry):
    block = {
        "type": "two_images_row",
        "images": [
            {"path": "missing_left.png", "width_cm": 7, "caption": "图1a"},
            {"path": "missing_right.png", "width_cm": 7, "caption": "图1b"},
        ],
    }
    registry.render(subdoc, block, style_map)
    assert len(subdoc.tables) == 1
    table = subdoc.tables[0]
    assert len(table.columns) == 2
    assert len(table.rows) == 1

def test_p1_blocks_in_registry():
    registry = create_default_registry()
    assert "rich_paragraph" in registry._renderers
    assert "note" in registry._renderers
    assert "quote" in registry._renderers
    assert "two_images_row" in registry._renderers


def test_appendix_table_block(subdoc, style_map, registry):
    block = {
        "type": "appendix_table",
        "title": "附表1：经费预算",
        "headers": ["项目", "金额"],
        "rows": [["设备费", "50"]],
    }
    registry.render(subdoc, block, style_map)
    assert len(subdoc.tables) == 1
    table = subdoc.tables[0]
    assert len(table.rows) == 2  # header + 1 data row


def test_checklist_block(subdoc, style_map, registry):
    block = {
        "type": "checklist",
        "items": [
            {"text": "已完成文献综述", "checked": True},
            {"text": "已提交伦理审查", "checked": False},
        ],
    }
    registry.render(subdoc, block, style_map)
    assert len(subdoc.paragraphs) == 2
    assert "☑" in subdoc.paragraphs[0].text
    assert "☐" in subdoc.paragraphs[1].text


def test_horizontal_rule_block(subdoc, style_map, registry):
    from docx.oxml.ns import qn
    block = {"type": "horizontal_rule"}
    registry.render(subdoc, block, style_map)
    assert len(subdoc.paragraphs) == 1
    p = subdoc.paragraphs[0]
    pPr = p._element.pPr
    assert pPr is not None
    pBdr = pPr.find(qn("w:pBdr"))
    assert pBdr is not None


def test_p2_blocks_in_registry():
    registry = create_default_registry()
    assert "appendix_table" in registry._renderers
    assert "checklist" in registry._renderers
    assert "horizontal_rule" in registry._renderers


def test_toc_placeholder_block(subdoc, style_map, registry):
    block = {"type": "toc_placeholder", "title": "目 录"}
    registry.render(subdoc, block, style_map)
    assert len(subdoc.paragraphs) >= 1


def test_code_block_single_line(subdoc, style_map, registry):
    block = {"type": "code_block", "code": "print('hello')", "language": "python"}
    registry.render(subdoc, block, style_map)
    assert len(subdoc.paragraphs) >= 1


def test_code_block_multiline(subdoc, style_map, registry):
    block = {"type": "code_block", "code": "def foo():\n    return 1\n\nfoo()"}
    registry.render(subdoc, block, style_map)
    assert len(subdoc.paragraphs) == 4


def test_formula_block_text_fallback(subdoc, style_map, registry):
    block = {"type": "formula", "latex": "E = mc^2", "caption": "公式1"}
    registry.render(subdoc, block, style_map)
    assert len(subdoc.paragraphs) >= 1


def test_columns_block(subdoc, style_map, registry):
    block = {
        "type": "columns",
        "count": 2,
        "columns": [
            [{"type": "paragraph", "text": "左列"}],
            [{"type": "paragraph", "text": "右列"}],
        ],
    }
    registry.render(subdoc, block, style_map)
    assert len(subdoc.tables) == 1
    table = subdoc.tables[0]
    assert len(table.columns) == 2


def test_p3_blocks_in_registry():
    registry = create_default_registry()
    assert "toc_placeholder" in registry._renderers
    assert "code_block" in registry._renderers
    assert "formula" in registry._renderers
    assert "columns" in registry._renderers
    assert "ascii_diagram" in registry._renderers


def test_ascii_diagram_block(subdoc, style_map, registry):
    block = {
        "type": "ascii_diagram",
        "ascii": "  +---+\n  | A |\n  +---+",
        "caption": "图1 架构示意",
        "font_size": 12,
    }
    registry.render(subdoc, block, style_map)
    # 验证有段落（图片段落 + caption + 空段落）
    assert len(subdoc.paragraphs) >= 2
    # 验证 caption 存在
    assert any("架构示意" in p.text for p in subdoc.paragraphs)


def test_style_map_override_heading(subdoc, style_map, registry):
    custom = dict(style_map)
    custom["heading_2"] = "Heading 3"
    registry.render(subdoc, {"type": "heading", "text": "标题", "level": 2}, custom)
    assert subdoc.paragraphs[0].style.name == "Heading 3"


def test_style_map_fallback_for_missing_style(subdoc, style_map, registry):
    custom = dict(style_map)
    custom["heading_2"] = "NonExistentStyle"
    registry.render(subdoc, {"type": "heading", "text": "标题", "level": 2}, custom)
    assert subdoc.paragraphs[0].style.name == "Heading 2"


def test_columns_block_nested_table(subdoc, style_map, registry):
    block = {
        "type": "columns",
        "count": 2,
        "columns": [
            [{"type": "table", "headers": ["A"], "rows": [["1"]]}],
            [{"type": "paragraph", "text": "右列"}],
        ],
    }
    registry.render(subdoc, block, style_map)
    assert len(subdoc.tables) == 1
    table = subdoc.tables[0]
    assert len(table.columns) == 2
    assert len(table.cell(0, 0).tables) == 1
    assert table.cell(0, 0).tables[0].rows[0].cells[0].text == "A"
    assert "右列" in table.cell(0, 1).text


def test_columns_block_nested_bullet_list(subdoc, style_map, registry):
    block = {
        "type": "columns",
        "count": 2,
        "columns": [
            [{"type": "bullet_list", "items": ["item1", "item2"]}],
            [{"type": "paragraph", "text": "右列"}],
        ],
    }
    registry.render(subdoc, block, style_map)
    table = subdoc.tables[0]
    left_cell = table.cell(0, 0)
    assert any("item1" in p.text for p in left_cell.paragraphs)
    assert any("item2" in p.text for p in left_cell.paragraphs)


def test_formula_block_matplotlib_image(subdoc, style_map, registry):
    pytest.importorskip("matplotlib")
    block = {"type": "formula", "latex": "E = mc^2", "caption": "公式1"}
    registry.render(subdoc, block, style_map)
    assert len(subdoc.paragraphs) >= 1


def test_formula_block_fallback_on_render_failure(subdoc, style_map, registry):
    pytest.importorskip("matplotlib")
    with patch("latex2mathml.converter.convert", side_effect=RuntimeError("mock")):
        block = {"type": "formula", "latex": "E = mc^2", "caption": "公式1"}
        registry.render(subdoc, block, style_map)
        assert len(subdoc.paragraphs) >= 1


def test_formula_block_omml_inserted(subdoc, style_map, registry):
    pytest.importorskip("latex2mathml")
    block = {"type": "formula", "latex": "E = mc^2", "caption": "公式1"}
    registry.render(subdoc, block, style_map)
    # 验证段落中包含 m:oMath 元素，且内部为 OMML 标签而非 MathML
    assert len(subdoc.paragraphs) >= 1
    p = subdoc.paragraphs[0]
    omath = p._element.find("{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath")
    assert omath is not None
    # 应包含 OMML 的 m:r / m:t / m:sSup，不应包含 MathML 的 math/mrow/mi
    mml_ns = "{http://www.w3.org/1998/Math/MathML}"
    assert omath.find(f"{mml_ns}math") is None
    assert omath.find(f"{mml_ns}mrow") is None
    assert omath.find(f"{mml_ns}mi") is None
    omml_ns = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
    assert omath.find(f"{omml_ns}r") is not None
    assert omath.find(f".//{omml_ns}t") is not None
    assert omath.find(f"{omml_ns}sSup") is not None


def test_columns_block_gap_cm(subdoc, style_map, registry):
    from docx.oxml.ns import qn
    block = {
        "type": "columns",
        "count": 2,
        "gap_cm": 1.0,
        "columns": [
            [{"type": "paragraph", "text": "左列"}],
            [{"type": "paragraph", "text": "右列"}],
        ],
    }
    registry.render(subdoc, block, style_map)
    assert len(subdoc.tables) == 1
    tbl_pr = subdoc.tables[0]._tbl.tblPr
    spacing = tbl_pr.find(qn("w:tblCellSpacing"))
    assert spacing is not None
    assert spacing.get(qn("w:w")) == "567"  # 1 cm = 567 twips
