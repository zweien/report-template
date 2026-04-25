"""生成覆盖全部 18 种 block 类型的测试模板。

用法:
    python scripts/build_test_template.py

输出:
    templates/test_all_blocks.docx
"""

from pathlib import Path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_cell_shading(cell, color: str):
    """设置单元格底纹颜色。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def add_paragraph_style(doc, name: str, *, font_name=None, font_size=None,
                        bold=None, italic=None, base_style=None):
    """添加段落样式，已存在则跳过。"""
    try:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        return
    if base_style:
        style.base_style = doc.styles[base_style]
    if font_name:
        style.font.name = font_name
    if font_size:
        style.font.size = font_size
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def add_table_style(doc, name: str):
    """添加表格样式，已存在则跳过。"""
    try:
        doc.styles.add_style(name, WD_STYLE_TYPE.TABLE)
    except ValueError:
        pass


def build_template(output_path: str):
    doc = Document()

    # ── 段落样式 ──────────────────────────────────────────────
    add_paragraph_style(doc, "Heading 1", bold=True)
    add_paragraph_style(doc, "Heading 2", bold=True)
    add_paragraph_style(doc, "Heading 3", bold=True)
    add_paragraph_style(doc, "Heading 4", bold=True)
    add_paragraph_style(doc, "Heading 5", bold=True)
    add_paragraph_style(doc, "Body Text")
    add_paragraph_style(doc, "Caption", italic=True)
    add_paragraph_style(doc, "Legend", italic=True)
    add_paragraph_style(doc, "Figure Paragraph")
    add_paragraph_style(doc, "List Bullet")
    add_paragraph_style(doc, "List Number")
    add_paragraph_style(doc, "Note", italic=True)
    add_paragraph_style(doc, "Quote", italic=True)
    add_paragraph_style(doc, "Checklist")
    add_paragraph_style(doc, "CodeBlock", font_name="Courier New")

    # ── 表格样式 ──────────────────────────────────────────────
    add_table_style(doc, "ResearchTable")
    add_table_style(doc, "AppendixTable")

    # ── 封面页 ────────────────────────────────────────────────
    doc.add_paragraph("{{PROJECT_NAME}}", style="Heading 2")
    doc.add_paragraph("申请单位：{{APPLICANT_ORG}}")
    doc.add_paragraph("项目负责人：{{PROJECT_LEADER}}")
    doc.add_paragraph("项目周期：{{PROJECT_PERIOD}}")
    doc.add_page_break()

    # ── 目录区 ────────────────────────────────────────────────
    doc.add_paragraph("{%p if ENABLE_TOC %}")
    doc.add_paragraph("目 录", style="Heading 1")
    doc.add_paragraph("{{p TOC_SUBDOC }}")
    doc.add_paragraph("{%p endif %}")
    doc.add_page_break()

    # ── 第一章：研究内容 ──────────────────────────────────────
    doc.add_paragraph("{%p if ENABLE_RESEARCH_CONTENT %}")
    doc.add_paragraph("一、研究内容与技术路线", style="Heading 1")
    doc.add_paragraph("{{p RESEARCH_CONTENT_SUBDOC }}")
    doc.add_paragraph("{%p endif %}")

    # ── 第二章：研究基础 ──────────────────────────────────────
    doc.add_paragraph("{%p if ENABLE_RESEARCH_BASIS %}")
    doc.add_paragraph("二、研究基础与条件保障", style="Heading 1")
    doc.add_paragraph("{{p RESEARCH_BASIS_SUBDOC }}")
    doc.add_paragraph("{%p endif %}")

    # ── 第三章：实施计划 ──────────────────────────────────────
    doc.add_paragraph("{%p if ENABLE_IMPLEMENTATION_PLAN %}")
    doc.add_paragraph("三、实施计划与进度安排", style="Heading 1")
    doc.add_paragraph("{{p IMPLEMENTATION_PLAN_SUBDOC }}")
    doc.add_paragraph("{%p endif %}")

    # ── 附件区 ────────────────────────────────────────────────
    doc.add_paragraph("{%p if ENABLE_APPENDICES %}")
    doc.add_paragraph("附件", style="Heading 1")
    doc.add_paragraph("{{p APPENDICES_SUBDOC }}")
    doc.add_paragraph("{%p endif %}")

    # ── 确保 OMML 命名空间已声明 ────────────────────────────────
    root = doc._element
    if "m" not in root.nsmap:
        root.set(
            "{http://www.w3.org/2000/xmlns/}m",
            "http://schemas.openxmlformats.org/officeDocument/2006/math",
        )

    # ── 保存 ──────────────────────────────────────────────────
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"模板已生成: {output_path}")


if __name__ == "__main__":
    build_template("templates/test_all_blocks.docx")
