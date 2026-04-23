from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional

from docxtpl import DocxTemplate
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm


DEFAULT_STYLE_MAP = {
    "heading_2": "Heading 2",
    "heading_3": "Heading 3",
    "body": "Body Text",
    "caption": "Caption",
    "legend": "Legend",
    "figure_paragraph": "Figure Paragraph",
    "table": "ResearchTable",
    "bullet_list": "List Bullet",
    "numbered_list": "List Number",
}


class GrantRendererError(Exception):
    pass


# -----------------------------
# Helpers
# -----------------------------

def _get_style_name(doc, preferred: str, fallback: str) -> str:
    style_names = {style.name for style in doc.styles}
    return preferred if preferred in style_names else fallback


def _set_table_borders(table) -> None:
    """兜底边框处理；如模板已定义完善表格样式，可删除此调用。"""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


# -----------------------------
# Block renderers
# -----------------------------

def add_heading_block(doc, text: str, level: int, style_map: Dict[str, str]) -> None:
    if level <= 2:
        style_name = _get_style_name(doc, style_map["heading_2"], "Heading 2")
    else:
        style_name = _get_style_name(doc, style_map["heading_3"], "Heading 3")
    doc.add_paragraph(text, style=style_name)


def add_paragraph_block(doc, text: str, style_map: Dict[str, str]) -> None:
    style_name = _get_style_name(doc, style_map["body"], "Normal")
    doc.add_paragraph(text, style=style_name)


def add_bullet_list_block(doc, items: List[str], style_map: Dict[str, str]) -> None:
    style_name = _get_style_name(doc, style_map["bullet_list"], style_map["body"])
    for item in items:
        p = doc.add_paragraph(style=style_name)
        p.add_run(str(item))


def add_numbered_list_block(doc, items: List[str], style_map: Dict[str, str]) -> None:
    style_name = _get_style_name(doc, style_map["numbered_list"], style_map["body"])
    for item in items:
        p = doc.add_paragraph(style=style_name)
        p.add_run(str(item))


def add_page_break_block(doc) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_table_block(doc, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
    if block.get("title"):
        caption_style = _get_style_name(doc, style_map["caption"], "Caption")
        doc.add_paragraph(str(block["title"]), style=caption_style)

    headers = block["headers"]
    rows = block["rows"]
    table_style = block.get("style") or style_map["table"]
    table_style = _get_style_name(doc, table_style, "Table Grid")

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = table_style

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = "" if header is None else str(header)

    for row in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = "" if value is None else str(value)

    if block.get("force_borders", True):
        _set_table_borders(table)

    body_style = _get_style_name(doc, style_map["body"], "Normal")
    doc.add_paragraph("", style=body_style)


def add_image_block(doc, block: Dict[str, Any], style_map: Dict[str, str]) -> None:
    image_path = Path(block["path"])

    figure_style = _get_style_name(doc, style_map["figure_paragraph"], style_map["body"])
    p = doc.add_paragraph(style=figure_style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if image_path.exists():
        run = p.add_run()
        width_cm = block.get("width_cm")
        if width_cm is not None:
            run.add_picture(str(image_path), width=Cm(float(width_cm)))
        else:
            run.add_picture(str(image_path))
    else:
        p.add_run(f"[图片缺失：{image_path}]")

    if block.get("caption"):
        caption_style = _get_style_name(doc, style_map["caption"], "Caption")
        cp = doc.add_paragraph(str(block["caption"]), style=caption_style)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if block.get("legend"):
        legend_style = _get_style_name(doc, style_map["legend"], style_map["body"])
        lp = doc.add_paragraph(str(block["legend"]), style=legend_style)
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    body_style = _get_style_name(doc, style_map["body"], "Normal")
    doc.add_paragraph("", style=body_style)


# -----------------------------
# Subdoc builders
# -----------------------------

def render_blocks(doc, blocks: Iterable[Dict[str, Any]], style_map: Dict[str, str]) -> None:
    for block in blocks:
        block_type = block["type"]

        if block_type == "heading":
            add_heading_block(doc, block["text"], block.get("level", 2), style_map)
        elif block_type == "paragraph":
            add_paragraph_block(doc, block["text"], style_map)
        elif block_type == "bullet_list":
            add_bullet_list_block(doc, block["items"], style_map)
        elif block_type == "numbered_list":
            add_numbered_list_block(doc, block["items"], style_map)
        elif block_type == "table":
            add_table_block(doc, block, style_map)
        elif block_type == "image":
            add_image_block(doc, block, style_map)
        elif block_type == "page_break":
            add_page_break_block(doc)
        else:
            raise GrantRendererError(f"Unsupported block type: {block_type}")


def build_subdoc(
    tpl: DocxTemplate,
    blocks: List[Dict[str, Any]],
    style_map: Dict[str, str],
    title: Optional[str] = None,
    title_level: int = 2,
) -> Any:
    subdoc = tpl.new_subdoc()
    if title:
        add_heading_block(subdoc, title, title_level, style_map)
    render_blocks(subdoc, blocks, style_map)
    return subdoc


# -----------------------------
# Context builders
# -----------------------------

def _normalize_style_map(payload: Dict[str, Any]) -> Dict[str, str]:
    style_map = dict(DEFAULT_STYLE_MAP)
    style_map.update(payload.get("style_map", {}))
    return style_map


def _build_sections_context(tpl: DocxTemplate, payload: Dict[str, Any], context: Dict[str, Any], style_map: Dict[str, str]) -> None:
    for section in payload.get("sections", []):
        section_id = section.get("id", "SECTION")
        placeholder = section.get("placeholder", f"{section_id.upper()}_SUBDOC")
        flag_name = section.get("flag_name", f"ENABLE_{section_id.upper()}")
        enabled = bool(section.get("enabled", True))

        context[flag_name] = enabled

        if enabled:
            subdoc = build_subdoc(
                tpl,
                section.get("blocks", []),
                style_map,
                title=section.get("subdoc_title"),
                title_level=section.get("subdoc_title_level", 2),
            )
            context[placeholder] = subdoc
        else:
            context[placeholder] = ""


def _build_individual_attachments_context(
    tpl: DocxTemplate, payload: Dict[str, Any], context: Dict[str, Any], style_map: Dict[str, str]
) -> List[Dict[str, Any]]:
    enabled_attachments: List[Dict[str, Any]] = []

    for attachment in payload.get("attachments", []):
        attachment_id = attachment.get("id", "APPENDIX")
        placeholder = attachment.get("placeholder", f"{attachment_id.upper()}_SUBDOC")
        flag_name = attachment.get("flag_name", f"ENABLE_{attachment_id.upper()}")
        enabled = bool(attachment.get("enabled", True))

        context[flag_name] = enabled

        if enabled:
            enabled_attachments.append(attachment)
            subdoc = build_subdoc(
                tpl,
                attachment.get("blocks", []),
                style_map,
                title=attachment.get("title"),
                title_level=attachment.get("title_level", 2),
            )
            context[placeholder] = subdoc
        else:
            context[placeholder] = ""

    return enabled_attachments


def _build_bundle_attachments_context(
    tpl: DocxTemplate,
    payload: Dict[str, Any],
    context: Dict[str, Any],
    style_map: Dict[str, str],
    enabled_attachments: List[Dict[str, Any]],
) -> None:
    bundle = payload.get("attachments_bundle", {})
    placeholder = bundle.get("placeholder", "APPENDICES_SUBDOC")
    flag_name = bundle.get("flag_name", "ENABLE_APPENDICES")
    bundle_enabled = bool(bundle.get("enabled", True)) and bool(enabled_attachments)
    page_break_between = bool(bundle.get("page_break_between_attachments", True))

    context[flag_name] = bundle_enabled

    if not bundle_enabled:
        context[placeholder] = ""
        return

    bundle_subdoc = tpl.new_subdoc()
    title_prefix_enabled = bool(bundle.get("include_attachment_title", True))

    for idx, attachment in enumerate(enabled_attachments):
        if idx > 0 and page_break_between:
            add_page_break_block(bundle_subdoc)

        attachment_title = attachment.get("title") if title_prefix_enabled else None
        if attachment_title:
            add_heading_block(
                bundle_subdoc,
                attachment_title,
                attachment.get("title_level", 2),
                style_map,
            )
        render_blocks(bundle_subdoc, attachment.get("blocks", []), style_map)

    context[placeholder] = bundle_subdoc


# -----------------------------
# Public API
# -----------------------------

def render_grant_advanced(template_path: str, output_path: str, payload: Dict[str, Any]) -> None:
    tpl = DocxTemplate(template_path)
    style_map = _normalize_style_map(payload)

    # 标量上下文：优先使用 payload["context"]，用于普通占位符
    context: Dict[str, Any] = dict(payload.get("context", {}))

    # 兼容旧版固定字段
    compatibility_map = {
        "PROJECT_NAME": payload.get("project_name"),
        "APPLICANT_ORG": payload.get("applicant_org"),
        "PROJECT_LEADER": payload.get("project_leader"),
    }
    for key, value in compatibility_map.items():
        if value is not None and key not in context:
            context[key] = value

    _build_sections_context(tpl, payload, context, style_map)
    enabled_attachments = _build_individual_attachments_context(tpl, payload, context, style_map)
    _build_bundle_attachments_context(tpl, payload, context, style_map, enabled_attachments)

    tpl.render(context, autoescape=True)
    tpl.save(output_path)


if __name__ == "__main__":
    payload_path = Path("grant_payload_advanced_demo.json")
    template_path = Path("grant_template_demo_clean_v3.docx")
    output_path = Path("grant_output_advanced_demo.docx")

    payload = json.loads(payload_path.read_text(encoding="utf-8"))
    render_grant_advanced(str(template_path), str(output_path), payload)
    print(f"Generated: {output_path.resolve()}")
