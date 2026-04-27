#!/usr/bin/env python3
"""生成带自定义格式的 report-engine 模板。"""

import argparse
import sys
from pathlib import Path

import docx.oxml
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── 样式定义 ──────────────────────────────────────────────────

PARAGRAPH_STYLES = [
    ("Heading 1", {"bold": True, "font_name": "黑体", "size": Pt(14),
                   "space_before": Pt(18), "space_after": Pt(0),
                   "line_spacing": Pt(18), "line_spacing_rule": WD_LINE_SPACING.AT_LEAST,
                   "alignment": WD_ALIGN_PARAGRAPH.LEFT}),
    ("Heading 2", {"bold": True, "font_name": "黑体", "size": Pt(10.5),
                   "space_before": Pt(18), "space_after": Pt(0),
                   "line_spacing": Pt(16), "line_spacing_rule": WD_LINE_SPACING.AT_LEAST,
                   "alignment": WD_ALIGN_PARAGRAPH.LEFT}),
    ("Heading 3", {"bold": False, "font_name": "宋体", "size": Pt(10.5),
                   "line_spacing": Pt(18), "line_spacing_rule": WD_LINE_SPACING.AT_LEAST,
                   "alignment": WD_ALIGN_PARAGRAPH.LEFT}),
    ("Heading 4", {"bold": False, "font_name": "宋体", "size": Pt(10.5),
                   "line_spacing": Pt(18), "line_spacing_rule": WD_LINE_SPACING.AT_LEAST,
                   "alignment": WD_ALIGN_PARAGRAPH.LEFT}),
    ("Heading 5", {"bold": False, "font_name": "宋体", "size": Pt(10.5),
                   "line_spacing": Pt(18), "line_spacing_rule": WD_LINE_SPACING.AT_LEAST,
                   "alignment": WD_ALIGN_PARAGRAPH.LEFT}),
    ("Body Text", {"bold": False, "font_name": "宋体", "size": Pt(10.5),
                   "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
                   "first_line_indent": Cm(0.76),
                   "space_before": Pt(0), "space_after": Pt(0),
                   "line_spacing": Pt(18), "line_spacing_rule": WD_LINE_SPACING.AT_LEAST}),
    ("TableCaption", {"bold": True, "font_name": "黑体", "size": Pt(10.5),
                       "alignment": WD_ALIGN_PARAGRAPH.CENTER,
                       "space_before": Pt(18), "space_after": Pt(0),
                       "line_spacing": Pt(18), "line_spacing_rule": WD_LINE_SPACING.AT_LEAST}),
    ("FigureCaption", {"bold": True, "font_name": "黑体", "size": Pt(10.5),
                        "alignment": WD_ALIGN_PARAGRAPH.CENTER,
                        "space_before": Pt(0), "space_after": Pt(18),
                        "line_spacing": Pt(18), "line_spacing_rule": WD_LINE_SPACING.AT_LEAST}),
    ("Legend", {"italic": True, "font_name": "宋体", "size": Pt(10.5)}),
    ("Figure Paragraph", {"font_name": "宋体", "size": Pt(10.5),
                           "alignment": WD_ALIGN_PARAGRAPH.CENTER}),
    ("List Bullet", {"font_name": "宋体", "size": Pt(10.5)}),
    ("List Number", {"font_name": "宋体", "size": Pt(10.5)}),
    ("Note", {"italic": True, "font_name": "宋体", "size": Pt(10.5)}),
    ("Quote", {"italic": True, "font_name": "宋体", "size": Pt(10.5)}),
    ("Checklist", {"font_name": "宋体", "size": Pt(10.5)}),
    ("CodeBlock", {"font_name": "Courier New", "size": Pt(10)}),
]

# 目录样式
TOC_STYLES = [
    ("TOC 1", {"font_name": "宋体", "size": Pt(10.5),
                "alignment": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
                "space_before": Pt(0), "space_after": Pt(0),
                "line_spacing": Pt(16), "line_spacing_rule": WD_LINE_SPACING.AT_LEAST}),
    ("TOC 2", {"font_name": "宋体", "size": Pt(10.5),
                "alignment": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
                "space_before": Pt(0), "space_after": Pt(0),
                "line_spacing": Pt(16), "line_spacing_rule": WD_LINE_SPACING.AT_LEAST}),
    ("TOC 3", {"font_name": "宋体", "size": Pt(10.5),
                "alignment": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
                "space_before": Pt(0), "space_after": Pt(0),
                "line_spacing": Pt(16), "line_spacing_rule": WD_LINE_SPACING.AT_LEAST}),
]

TABLE_STYLES = ["ResearchTable", "AppendixTable"]

BLACK = RGBColor(0x00, 0x00, 0x00)

SCALAR_PLACEHOLDERS = [
    ("项目名称", "（请填写项目名称）"),
    ("申请单位", "（请填写申请单位）"),
    ("项目负责人", "（请填写项目负责人）"),
    ("项目周期", "（请填写项目周期）"),
]

SECTIONS = [
    "立项依据",
    "研究目标",
    "拟解决的关键科学问题",
    "研究内容",
    "与其研究成果和应用前景",
    "研究团队研究基础及保障条件",
    "项目组成员",
    "经费概算",
]


def _set_font(style, font_name: str):
    """设置样式的字体，并清除 theme 属性避免 Word 优先使用主题字体。"""
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = docx.oxml.OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    # 移除所有 theme 属性，否则 Word 会优先使用主题字体
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        if rFonts.get(qn(attr)) is not None:
            del rFonts.attrib[qn(attr)]
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    style.font.name = font_name


def add_styles(doc: Document):
    """添加并配置所有必需样式。"""
    for name, opts in PARAGRAPH_STYLES:
        try:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            style = doc.styles[name]

        font = style.font
        pf = style.paragraph_format

        if "font_name" in opts:
            _set_font(style, opts["font_name"])
        if "size" in opts:
            font.size = opts["size"]
        if "bold" in opts:
            font.bold = opts["bold"]
        if "italic" in opts:
            font.italic = opts["italic"]

        font.color.rgb = BLACK

        if "alignment" in opts:
            pf.alignment = opts["alignment"]
        if "space_before" in opts:
            pf.space_before = opts["space_before"]
        if "space_after" in opts:
            pf.space_after = opts["space_after"]
        if "first_line_indent" in opts:
            pf.first_line_indent = opts["first_line_indent"]
        if "line_spacing" in opts:
            pf.line_spacing = opts["line_spacing"]
        if "line_spacing_rule" in opts:
            pf.line_spacing_rule = opts["line_spacing_rule"]

    # 目录样式
    for name, opts in TOC_STYLES:
        try:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            style = doc.styles[name]

        font = style.font
        pf = style.paragraph_format

        if "font_name" in opts:
            _set_font(style, opts["font_name"])
        if "size" in opts:
            font.size = opts["size"]

        font.color.rgb = BLACK

        if "alignment" in opts:
            pf.alignment = opts["alignment"]
        if "space_before" in opts:
            pf.space_before = opts["space_before"]
        if "space_after" in opts:
            pf.space_after = opts["space_after"]
        if "line_spacing" in opts:
            pf.line_spacing = opts["line_spacing"]
        if "line_spacing_rule" in opts:
            pf.line_spacing_rule = opts["line_spacing_rule"]

    # 表格样式（仅创建，python-docx 对表格样式的格式支持有限）
    for name in TABLE_STYLES:
        try:
            doc.styles.add_style(name, WD_STYLE_TYPE.TABLE)
        except ValueError:
            pass


def add_cover(doc: Document):
    """添加封面页。"""
    doc.add_paragraph("{{项目名称}}", style="Heading 1")
    doc.add_paragraph("申请单位：{{申请单位}}", style="Body Text")
    doc.add_paragraph("项目负责人：{{项目负责人}}", style="Body Text")
    doc.add_paragraph("项目周期：{{项目周期}}", style="Body Text")
    doc.add_page_break()


def add_toc(doc: Document):
    """添加目录区。"""
    doc.add_paragraph("{%p if ENABLE_目录 %}")
    doc.add_paragraph("{{p 目录_SUBDOC }}")
    doc.add_paragraph("{%p endif %}")
    doc.add_page_break()


def add_section(doc: Document, title: str, prefix: str, number: int):
    """添加一个带条件开关的章节。"""
    chinese_nums = "一二三四五六七八九十"
    num_str = chinese_nums[number - 1] if number <= len(chinese_nums) else str(number)
    # 将中文标题中的特殊字符替换为下划线，生成合法的 Jinja 标识符
    safe_prefix = prefix.replace(" ", "_").replace("、", "_").replace("/", "_")
    doc.add_paragraph(f"{{%p if ENABLE_{safe_prefix} %}}")
    doc.add_paragraph(f"{num_str}、{title}", style="Heading 1")
    doc.add_paragraph(f"{{{{p {safe_prefix}_SUBDOC }}}}")
    doc.add_paragraph("{%p endif %}")


def build_template(output_path: str):
    """生成完整模板。"""
    doc = Document()
    add_styles(doc)
    add_cover(doc)
    add_toc(doc)

    for i, title in enumerate(SECTIONS, 1):
        add_section(doc, title, title, i)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"模板已生成: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="生成带自定义格式的 report-engine 模板")
    parser.add_argument("--output", required=True, help="输出路径")
    args = parser.parse_args()
    build_template(args.output)


if __name__ == "__main__":
    main()
