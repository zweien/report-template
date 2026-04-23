from pathlib import Path
from typing import Any, Dict, List

from docxtpl import DocxTemplate
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm


def _set_table_borders(table) -> None:
    """
    可选的兜底边框处理。
    如果模板里的 ResearchTable 样式已经定义好了边框，可删除此函数调用。
    """
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_heading_block(doc, text: str, level: int) -> None:
    style_name = "Heading 2" if level == 2 else "Heading 3"
    doc.add_paragraph(text, style=style_name)


def add_paragraph_block(doc, text: str) -> None:
    doc.add_paragraph(text, style="Body Text")


def add_bullet_list_block(doc, items: List[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="Body Text")
        p.style = "List Bullet" if "List Bullet" in [s.name for s in doc.styles] else "Body Text"
        p.add_run(item)


def add_table_block(doc, block: Dict[str, Any]) -> None:
    if block.get("title"):
        doc.add_paragraph(block["title"], style="Caption")

    headers = block["headers"]
    rows = block["rows"]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = block.get("style", "ResearchTable")

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)

    for row in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = "" if value is None else str(value)

    _set_table_borders(table)
    doc.add_paragraph("", style="Body Text")


def add_image_block(doc, block: Dict[str, Any]) -> None:
    image_path = Path(block["path"])

    p = doc.add_paragraph(style="Figure Paragraph")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if image_path.exists():
        run = p.add_run()
        run.add_picture(str(image_path), width=Cm(block.get("width_cm", 14)))
    else:
        placeholder = f"[图片缺失：{image_path}]"
        p.add_run(placeholder)

    if block.get("caption"):
        cp = doc.add_paragraph(block["caption"], style="Caption")
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if block.get("legend"):
        lp = doc.add_paragraph(block["legend"], style="Legend")
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph("", style="Body Text")


def build_research_subdoc(tpl: DocxTemplate, blocks: List[Dict[str, Any]]):
    subdoc = tpl.new_subdoc()

    for block in blocks:
        block_type = block["type"]

        if block_type == "heading":
            add_heading_block(subdoc, block["text"], block.get("level", 2))
        elif block_type == "paragraph":
            add_paragraph_block(subdoc, block["text"])
        elif block_type == "bullet_list":
            add_bullet_list_block(subdoc, block["items"])
        elif block_type == "table":
            add_table_block(subdoc, block)
        elif block_type == "image":
            add_image_block(subdoc, block)
        else:
            raise ValueError(f"Unsupported block type: {block_type}")

    return subdoc


def render_grant(template_path: str, output_path: str, payload: Dict[str, Any]) -> None:
    tpl = DocxTemplate(template_path)

    research_subdoc = build_research_subdoc(tpl, payload["research_content"])

    context = {
        "PROJECT_NAME": payload.get("project_name", ""),
        "APPLICANT_ORG": payload.get("applicant_org", ""),
        "PROJECT_LEADER": payload.get("project_leader", ""),
        "RESEARCH_CONTENT_SUBDOC": research_subdoc,
    }

    tpl.render(context, autoescape=True)
    tpl.save(output_path)


if __name__ == "__main__":
    import json

    payload_path = Path("grant_payload_demo.json")
    template_path = Path("grant_template_demo_clean_v3.docx")
    output_path = Path("grant_output_demo.docx")

    payload = json.loads(payload_path.read_text(encoding="utf-8"))
    render_grant(str(template_path), str(output_path), payload)
    print(f"Generated: {output_path.resolve()}")
